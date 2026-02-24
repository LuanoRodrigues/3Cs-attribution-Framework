import html
import json
import os
import subprocess
# src/ui/bibliographic_analysis_widget.py
import sys
import re
import io
from pathlib import Path
import pandas as pd
import logging
import traceback

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QComboBox,
    QFileDialog, QMessageBox, QStackedWidget, QTableWidget, QTableWidgetItem,
    QHeaderView, QAbstractItemView, QListWidget, QListWidgetItem, QRadioButton,
    QButtonGroup, QSizePolicy, QFrame, QTextBrowser, QApplication, QLineEdit, QGroupBox,
    QSpacerItem, QCheckBox
)
from PyQt6.QtCore import Qt, pyqtSignal, QThread, QObject, pyqtSlot, QWaitCondition, QMutex, QDateTime, QUrl, \
    QTemporaryFile, QDir, QTimer
from PyQt6.QtGui import QIcon, QImage, QPixmap, QColor, QPalette

from bibliometric_analysis_tool.utils.data_processing import load_data_from_source_for_widget
from gpt_api import _process_batch_for

# --- WebEngine Handling ---
WEBENGINE_AVAILABLE = False
try:
    from PyQt6.QtWebEngineWidgets import QWebEngineView
    from PyQt6.QtWebEngineCore import QWebEngineSettings

    WEBENGINE_AVAILABLE = True
    logging.info("PyQtWebEngine found and will be used for interactive plots and full report display.")
except ImportError:
    logging.warning(
        "PyQtWebEngine not found. Interactive Plotly plots and full HTML report view will be limited. Install PyQt6-WebEngine.")


    class QWebEngineView(QTextBrowser):  # type: ignore
        def __init__(self, parent=None):
            super().__init__(parent)
            self.setOpenExternalLinks(True)
            self.setHtml(
                "<div style='padding:20px; text-align:center;'><b>Interactive plots & Rich HTML Reports require PyQtWebEngine.</b><br>"
                "Please install it (e.g., <code>pip install PyQt6-WebEngine</code>).<br>"
                "Fallback display will be used.</div>")

        def settings(self): return None

        def page(self): return self

        def runJavaScript(self, s, callback=None):
            if callback: callback(None)

        def setHtml(self, html, baseUrl=None):  # type: ignore
            super().setHtml(html)

# --- Corrected Core Util Imports ---
from src.core.utils.constants import (
    STEP_LOAD_DATA, STEP_ANALYZE_AUTHORS, STEP_ANALYZE_KEYWORDS,
    STEP_ANALYZE_CITATIONS, STEP_GENERATE_COAUTHORSHIP_NET,
    STEP_GENERATE_KEYWORD_COOCCURRENCE_NET,  # STEP_GENERATE_CITATION_NET is often conceptual
    STEP_GENERATE_TIMELINE,
    STEP_AI_ABSTRACT, STEP_AI_INTRODUCTION, STEP_AI_LITERATURE_REVIEW_SUMMARY,
    STEP_AI_METHODOLOGY_REVIEW, STEP_AI_DISCUSSION, STEP_AI_CONCLUSION,
    STEP_AI_LIMITATIONS, STEP_ASSEMBLE_REPORT_HTML, STEP_AI_KEYWORD_GRAPH_ANALYSIS, STEP_AI_AUTHOR_GRAPH_ANALYSIS,
    STEP_AI_TIMELINE_ANALYSIS, STEP_AI_AUTHOR_FOCUS_KEYWORDS, STEP_AI_SUGGEST_CODING_KEYWORDS, STEP_AI_SOURCES_ANALYSIS,
    STEP_AI_COAUTHORSHIP_NETWORK_ANALYSIS, STEP_AI_GENERAL_KEYWORD_THEMES_ANALYSIS, STEP_AI_CITATION_TABLE_ANALYSIS,
    STEP_AI_FUNDERS_ANALYSIS, STEP_AI_PDF_CONTENT_SUMMARY, STEP_AI_AFFILIATIONS_ANALYSIS, STEP_AI_COUNTRIES_ANALYSIS,
    STEP_AI_DOC_TYPES_ANALYSIS, STEP_AI_KEYWORD_CLUSTER_SEARCH_TERMS, STEP_EXTRACT_PDF_CONTENT_FOR_KEYWORDS,
    ## MODIFICATION: Add new step constants for the chronological review
    STEP_ANALYZE_TIMELINE_PERIODS, STEP_AI_GENERATE_CHRONOLOGICAL_NARRATIVE,
)
from src.core.utils.data_processing import (
    # load_data_from_source_for_widget,
    analyze_authors_detailed_plotly,
    analyze_keywords_detailed_plotly,
    analyze_citations_data_plotly,
    generate_network_graph_detailed_plotly,
    generate_timeline_detailed_plotly, analyze_publication_trends, describe_trend_data,
    ## MODIFICATION: Import the new period analysis function
    analyze_and_visualize_timeline_periods,
)
from src.core.utils.ai_services import (
    generate_ai_abstract, generate_ai_introduction,
    generate_ai_literature_review_summary,
    generate_ai_methodology_insights,
    generate_ai_discussion, generate_ai_conclusion,
    generate_ai_limitations,
    generate_thematic_keywords_for_pdf_search, # Utility
    generate_ai_author_focus_keywords,      # Utility
    generate_ai_keyword_cluster_search_terms, # Utility
    # Orchestrators for existing complex analyses
    generate_ai_author_graph_analysis,    # Orchestrator
    generate_ai_keyword_graph_analysis,   # Orchestrator
    generate_ai_timeline_analysis,        # Orchestrator (ensure this is now an orchestrator)
    # New orchestrators for other figures/tables - you will define these in ai_services.py
    generate_ai_doc_types_analysis,
    generate_ai_affiliations_analysis,
    generate_ai_countries_analysis,
    generate_ai_general_keyword_themes_analysis,
    generate_ai_coauthorship_network_analysis, # Distinct from author productivity graph analysis
    generate_ai_sources_analysis,
    generate_ai_citation_table_analysis,
    generate_ai_funders_analysis,
    generate_ai_pdf_content_summary, # General summary of notes
    ## MODIFICATION: Import the new AI narrative generation function
    generate_ai_chronological_narrative,
)
from src.core.utils.report_generation import (
    assemble_report_html
)

# If pdf_processing functions are directly called by worker lambdas (they might be through ai_services or data_processing)
# from src.core.utils.pdf_processing import extract_content_for_keywords
# ─── Shared CSS for all QTextBrowser / QWebEngine HTML ─────────────────────────
# COMMON_MONITOR_CSS = """
# <style>
#   :root {
#       --fg-light: #1e1e1e;   --bg-light: #ffffff;
#       --fg-dark : #dcdcdc;   --bg-dark : #1e1e1e;
#   }
#   body        { margin:0; padding:0;
#                 color:var(--fg-light); background:var(--bg-light);
#                 font: 14px/1.5 "Segoe UI", Tahoma, sans-serif; }
#   html[data-dark="1"] body {
#                 color:var(--fg-dark); background:var(--bg-dark); }
#   h1,h2,h3,h4 { margin: 0.6em 0 0.4em; font-weight:600; }
#   pre, code   { background:transparent; color:inherit;
#                 font: 13px "Fira Code","Consolas",monospace;
#                 white-space:pre-wrap; border:none; }
#   details > summary { cursor:pointer; }
#   ul          { margin-top:0.2em; }
# </style>
# """
COMMON_MONITOR_CSS = """
<style>
  :root {
      /* Base Colors */
      --fg-light: #1e1e1e;   --bg-light: #ffffff;
      --fg-dark: #dcdcdc;    --bg-dark: #1e1e1e;

      /* Dark Theme - PyCharm-like */
      --monitor-bg: #1e1e1e;
      --monitor-fg: #a9b7c6;
      --monitor-border: #3a3a3a;
      --monitor-accent: #4e94ce;
      --monitor-warning: #d19a66;
      --monitor-error: #d1676a;
      --monitor-success: #98c379;
      --monitor-info: #56b6c2;

      /* Code/Log Specific */
      --log-debug: #5c6370;
      --log-info: var(--monitor-info);
      --log-warning: var(--monitor-warning);
      --log-error: var(--monitor-error);
      --log-timestamp: #6a9955;

      /* UI Elements */
      --table-header-bg: #2b2b2b;
      --table-row-alt: #252526;
      --table-hover: #37373d;
      --link-color: #4e94ce;
  }

  body {
      margin: 0;
      padding: 15px;
      color: var(--monitor-fg);
      background-color: var(--monitor-bg);
      font-family: 'Consolas', 'Monaco', 'JetBrains Mono', monospace;
      font-size: 13px;
      line-height: 1.5;
  }

  /* Typography */
  h1, h2, h3, h4 {
      margin: 1.2em 0 0.6em;
      font-weight: 600;
      line-height: 1.3;
      color: var(--monitor-fg);
      border-bottom: 1px solid var(--monitor-border);
      padding-bottom: 0.3em;
  }
  h1 { font-size: 1.8em; color: #dcdcdc; }
  h2 { font-size: 1.5em; color: #d7ba7d; }
  h3 { font-size: 1.25em; }
  h4 { font-size: 1.1em; border-bottom: none; color: #9cdcfe; }

  /* Log Messages */
  .log-entry {
      margin-bottom: 4px;
      font-family: monospace;
      white-space: pre-wrap;
  }
  .log-timestamp {
      color: var(--log-timestamp);
      font-weight: bold;
  }
  .log-debug { color: var(--log-debug); }
  .log-info { color: var(--log-info); }
  .log-warning { color: var(--log-warning); }
  .log-error { color: var(--log-error); }
  .log-success { color: var(--monitor-success); }

  /* Code Blocks and Preformatted Text */
  pre, code {
      background-color: #282c34;
      color: #abb2bf;
      border: 1px solid var(--monitor-border);
      border-radius: 4px;
      padding: 12px;
      overflow-x: auto;
      font-family: 'Consolas', 'Monaco', 'JetBrains Mono', monospace;
      line-height: 1.4;
      tab-size: 4;
  }

  /* Lists */
  ul, ol {
      margin-top: 0.5em;
      padding-left: 25px;
      color: var(--monitor-fg);
  }
  li {
      margin-bottom: 0.4em;
      line-height: 1.5;
  }
  ul li::marker {
      color: #d19a66;
  }
  ol li::marker {
      color: #56b6c2;
  }

  /* Tables */
  table.dataframe {
      border-collapse: collapse;
      width: 100%;
      margin: 1.5em 0;
      border: 1px solid var(--monitor-border);
      font-size: 0.9em;
  }
  table.dataframe th {
      background-color: var(--table-header-bg);
      color: var(--monitor-fg);
      font-weight: 600;
      position: sticky;
      top: 0;
      padding: 10px 12px;
      text-align: left;
      border: 1px solid var(--monitor-border);
  }
  table.dataframe td {
      padding: 8px 12px;
      border: 1px solid var(--monitor-border);
      vertical-align: top;
  }
  table.dataframe tr:nth-child(even) {
      background-color: var(--table-row-alt);
  }
  table.dataframe tr:hover {
      background-color: var(--table-hover);
  }

  /* Text Elements */
  .text-summary {
      white-space: pre-wrap;
      line-height: 1.6;
      padding: 12px;
      background-color: #282c34;
      border-left: 4px solid #4e94ce;
      margin: 1em 0;
  }
  .text-summary p {
      margin: 0.5em 0;
  }

  /* Links */
  a {
      color: var(--link-color);
      text-decoration: none;
  }
  a:hover {
      text-decoration: underline;
  }

  /* Images */
  img {
      max-width: 95%;
      height: auto;
      display: block;
      margin: 20px auto;
      border: 1px solid var(--monitor-border);
      border-radius: 4px;
      box-shadow: 0 2px 8px rgba(0,0,0,0.3);
  }

  /* Utility Classes */
  .placeholder-text {
      color: #7f8c8d;
      font-style: italic;
      padding: 20px;
      text-align: center;
  }
  .status-box {
      padding: 12px;
      margin: 1em 0;
      border-radius: 4px;
      border-left: 4px solid;
  }
  .status-info {
      background-color: rgba(86, 182, 194, 0.1);
      border-left-color: var(--monitor-info);
  }
  .status-warning {
      background-color: rgba(209, 154, 102, 0.1);
      border-left-color: var(--monitor-warning);
  }
  .status-error {
      background-color: rgba(209, 103, 106, 0.1);
      border-left-color: var(--monitor-error);
  }
</style>
"""

# --- Analysis Configuration (Module-level constants for names) ---
ANALYSIS_TYPE_AUTHOR = "Author Analysis (Single)"
ANALYSIS_TYPE_KEYWORD = "Keyword Analysis (Single)"
ANALYSIS_TYPE_TIMELINE = "Publication Timeline (Single)"
ANALYSIS_TYPE_COAUTHOR_NET = "Co-authorship Network (Single)"
ANALYSIS_TYPE_FULL_REVIEW = "Full Bibliographic Review (Research & Writing)"
## MODIFICATION: Define the new analysis type constant
ANALYSIS_TYPE_CHRONOLOGICAL_REVIEW = "Chronological Literature Review"

PHASE_RESEARCH = "Research Phase: Data Analysis & Visualization"
PHASE_WRITING = "Writing Phase: AI-Assisted Report Generation"
## MODIFICATION: Add new phase constants for chronological review clarity
PHASE_CHRONO_ANALYSIS = "Phase 1: Chronological Analysis"
PHASE_CHRONO_WRITING = "Phase 2: AI Narrative Generation"
PHASE_CHRONO_SYNTHESIS = "Phase 3: Final Report Assembly"


# AI Model Configuration Data
AI_MODELS_CONFIG = {
    "OpenAI": {
        "gpt-4.1-mini-2025-04-14": "gpt-4o-mini",
        "GPT-4o": "gpt-4o",
        "GPT-4 Turbo": "gpt-4-turbo",

    },
    "DeepSeek": {
        "DeepSeek-Chat": "deepseek-chat",
        "DeepSeek-Coder": "deepseek-coder",
    },
}
def _simple_list_as_html(label: str, items: list[str]) -> str:
    if not items:
        return ""
    li = "".join(f"<li>{html.escape(i)}</li>" for i in items)
    return f"<h4>{html.escape(label)}</h4><ul>{li}</ul>"

class BibliographicAnalysisWorker(QObject):
    step_started = pyqtSignal(str, str)
    log_message_ready = pyqtSignal(str)
    data_summary_ready = pyqtSignal(object, str)
    table_data_ready = pyqtSignal(object, str)
    plotly_graph_ready = pyqtSignal(object, str)
    html_section_ready = pyqtSignal(object, str)
    step_completed = pyqtSignal(str, str, bool, object)
    analysis_finished = pyqtSignal(bool, str)
    paused_for_manual_step = pyqtSignal(str, str)

    live_table_ready = pyqtSignal(object, str)  # Emits: pandas.DataFrame, title_suffix
    live_plotly_ready = pyqtSignal(object, str)  # Emits: Plotly Figure object, title_suffix
    live_html_ready = pyqtSignal(object, str)  # Emits: HTML string or dict (for AI steps), title_suffix
    live_image_ready = pyqtSignal(str, str)  # Emits: string path to image, title_suffix

    def __init__(self, data_source_type, data_file_path, zotero_collection_name,
                 research_topic,  # NEW PARAMETER
                 analysis_type, is_manual_mode, ai_provider_key,
                 ai_model_api_name, current_theme_name, store_only: bool = False,
        read: bool = False):
        super().__init__()

        # ── BASIC STATE ──────────────────────────────────────────────────────────
        self.research_topic = research_topic
        self.data_source_type = data_source_type
        self.data_file_path = data_file_path
        self.zotero_collection_name = zotero_collection_name
        self.analysis_type = analysis_type
        self.is_manual_mode = is_manual_mode
        self.ai_provider_key = ai_provider_key
        self.ai_model_api_name = ai_model_api_name
        self.current_theme_name = current_theme_name
        self.store_only = bool(store_only)
        self.read = bool(read)


        self._bibliographic_data_df = None
        self._results_so_far = {}

        # ── CONTROL FLAGS / SYNC PRIMITIVES ─────────────────────────────────────
        self._is_running = False
        self._is_paused = False
        self._mutex = QMutex()
        self._pause_condition = QWaitCondition()

        # buffer that collects log lines belonging to the *current* step
        self._current_step_logs_buffer: list[str] = []

        # ── LIVE-STREAM LOG HANDLER  (safe even after QObject destruction) ──────
        import weakref, logging
        if hasattr(self, "_qt_log_handler"):  # remove any old one
            logging.getLogger().removeHandler(self._qt_log_handler)

        class _QtSignalLogHandler(logging.Handler):
            """
            Forwards each std-library logging record to outer.log_message_ready
            *while* the outer QObject is still alive.  If the worker is destroyed
            later, any further log record silently detaches this handler, avoiding
            RuntimeError (‘wrapped C/C++ object … has been deleted’).
            """

            def __init__(self, outer: QObject):
                super().__init__()
                self._outer_ref = weakref.ref(outer)  # avoid ref-cycle

            def emit(self, record):
                outer = self._outer_ref()
                if outer is None:  # QObject GC’ed
                    logging.getLogger().removeHandler(self)
                    return
                try:
                    msg = self.format(record)
                except Exception:
                    msg = record.getMessage()
                try:
                    outer.log_message_ready.emit(msg)
                except RuntimeError:  # C++ object gone
                    logging.getLogger().removeHandler(self)

        self._qt_log_handler = _QtSignalLogHandler(self)
        self._qt_log_handler.setFormatter(logging.Formatter("%(message)s"))
        _root_logger = logging.getLogger()
        _root_logger.addHandler(self._qt_log_handler)
        _root_logger.setLevel(logging.INFO)

        # auto-detach when this worker gets deleted
        self.destroyed.connect(lambda: _root_logger.removeHandler(self._qt_log_handler))


        ## MODIFICATION: Define the new step configuration for the chronological literature review
        self.chronological_review_steps_config = {
            PHASE_CHRONO_ANALYSIS: [
                (STEP_LOAD_DATA, load_data_from_source_for_widget, "raw_df_full_summary", True, True),
                (STEP_GENERATE_TIMELINE,
                 lambda df_arg, cb_arg, res_sf_arg: generate_timeline_detailed_plotly(df_arg, cb_arg),
                 "plotly_graph", True, True),
                (STEP_ANALYZE_TIMELINE_PERIODS,
                 lambda df_arg, cb_arg, res_sf_arg: analyze_and_visualize_timeline_periods(df_arg, cb_arg),
                 "period_analysis_package", True, True),
            ],
            PHASE_CHRONO_WRITING: [
                (STEP_AI_GENERATE_CHRONOLOGICAL_NARRATIVE,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_chronological_narrative(df_arg, res_sf_arg, cb_arg,
                                                                                       self.ai_provider_key,
                                                                                       self.ai_model_api_name),
                 "html_section", True, True),
            ],
            PHASE_CHRONO_SYNTHESIS: [
                (STEP_AI_METHODOLOGY_REVIEW,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_methodology_insights(df_arg, res_sf_arg, cb_arg,
                                                                                     self.ai_provider_key,
                                                                                     self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_DISCUSSION,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_discussion(df_arg, res_sf_arg, cb_arg,
                                                                           self.ai_provider_key,
                                                                           self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_LIMITATIONS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_limitations(df_arg, res_sf_arg, cb_arg,
                                                                            self.ai_provider_key,
                                                                            self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_CONCLUSION,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_conclusion(df_arg, res_sf_arg, cb_arg,
                                                                           self.ai_provider_key,
                                                                           self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_INTRODUCTION,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_introduction(df_arg, res_sf_arg, cb_arg,
                                                                             self.ai_provider_key,
                                                                             self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_ABSTRACT, lambda df_arg, cb_arg, res_sf_arg: generate_ai_abstract(df_arg, res_sf_arg, cb_arg,
                                                                                           self.ai_provider_key,
                                                                                           self.ai_model_api_name),
                 "html_section", True, True),
                (STEP_ASSEMBLE_REPORT_HTML,
                 lambda df_arg, cb_arg, res_sf_arg: assemble_report_html(self._bibliographic_data_df, res_sf_arg,
                                                                         cb_arg, theme_name=self.current_theme_name),
                 "full_html_report", True, True),
            ]
        }


        self.full_review_steps_config = {
            PHASE_RESEARCH: [
                (STEP_LOAD_DATA, load_data_from_source_for_widget, "raw_df_full_summary", True, True),
                (STEP_ANALYZE_AUTHORS,
                 lambda df_arg, cb_arg, res_sf_arg: analyze_authors_detailed_plotly(df_arg, cb_arg),
                 "plotly_graph_df", True, True),
                (STEP_ANALYZE_KEYWORDS,
                 lambda df_arg, cb_arg, res_sf_arg: analyze_keywords_detailed_plotly(df_arg, cb_arg),
                 "plotly_graph_df", True, True),
                (STEP_ANALYZE_CITATIONS,
                 lambda df_arg, cb_arg, res_sf_arg: analyze_citations_data_plotly(df_arg, cb_arg),
                 "plotly_graph_df", True, True),
                (STEP_GENERATE_TIMELINE,
                 lambda df_arg, cb_arg, res_sf_arg: generate_timeline_detailed_plotly(df_arg, cb_arg),
                 "plotly_graph", True, True),
                (STEP_GENERATE_COAUTHORSHIP_NET,
                 lambda df_arg, cb_arg, res_sf_arg: generate_network_graph_detailed_plotly(df_arg, "authors", cb_arg),
                 "plotly_graph_dict", False, True),
                (STEP_GENERATE_KEYWORD_COOCCURRENCE_NET,
                 lambda df_arg, cb_arg, res_sf_arg: generate_network_graph_detailed_plotly(df_arg, "keywords", cb_arg),
                 "plotly_graph_dict", True, True),
                (STEP_AI_SUGGEST_CODING_KEYWORDS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_thematic_keywords_for_pdf_search(
                     df_arg, res_sf_arg.get("research_question", "N/A"),
                     self.ai_provider_key, self.ai_model_api_name, cb_arg, res_sf_arg),
                 "keyword_list", False, True),
                (STEP_AI_AUTHOR_FOCUS_KEYWORDS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_author_focus_keywords(
                     res_sf_arg.get(STEP_ANALYZE_AUTHORS, {}).get("raw_df_table", pd.DataFrame()),
                     df_arg,
                     self.ai_provider_key, self.ai_model_api_name, cb_arg, res_sf_arg),
                 "author_keywords_map", False, True),
                (STEP_AI_KEYWORD_CLUSTER_SEARCH_TERMS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_keyword_cluster_search_terms(
                     ((res_sf_arg or {}).get(STEP_LOAD_DATA, {}) or {}).get("collection_name_for_title", "topic"),
                     (((res_sf_arg or {}).get(f"{STEP_GENERATE_KEYWORD_COOCCURRENCE_NET}_clusters_info") or {})
                      .get("summary_text") or "No cluster summary."),
                     df_arg.head(3) if isinstance(df_arg, pd.DataFrame) else pd.DataFrame(),
                     self.ai_provider_key, self.ai_model_api_name, cb_arg, (res_sf_arg or {})
                 ), "keyword_cluster_search_terms_map", False, True),
            ],
            PHASE_WRITING: [
                (STEP_AI_METHODOLOGY_REVIEW,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_methodology_insights(df_arg, res_sf_arg, cb_arg,
                                                                                     self.ai_provider_key,
                                                                                     self.ai_model_api_name),
                 "html_section", True, True),
                (STEP_AI_TIMELINE_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_timeline_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                  self.ai_provider_key,
                                                                                  self.ai_model_api_name),
                 "html_section", True, True),
                (STEP_AI_DOC_TYPES_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_doc_types_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                   self.ai_provider_key,
                                                                                   self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_AUTHOR_GRAPH_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_author_graph_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                      self.ai_provider_key,
                                                                                      self.ai_model_api_name),
                 "html_section", True, True),
                (STEP_AI_AFFILIATIONS_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_affiliations_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                      self.ai_provider_key,
                                                                                      self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_COUNTRIES_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_countries_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                   self.ai_provider_key,
                                                                                   self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_GENERAL_KEYWORD_THEMES_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_general_keyword_themes_analysis(df_arg, res_sf_arg,
                                                                                                cb_arg,
                                                                                                self.ai_provider_key,
                                                                                                self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_COAUTHORSHIP_NETWORK_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_coauthorship_network_analysis(df_arg, res_sf_arg,
                                                                                              cb_arg,
                                                                                              self.ai_provider_key,
                                                                                              self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_KEYWORD_GRAPH_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_keyword_graph_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                       self.ai_provider_key,
                                                                                       self.ai_model_api_name),
                 "html_section", True, True),
                (STEP_AI_SOURCES_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_sources_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                 self.ai_provider_key,
                                                                                 self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_CITATION_TABLE_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_citation_table_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                        self.ai_provider_key,
                                                                                        self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_FUNDERS_ANALYSIS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_funders_analysis(df_arg, res_sf_arg, cb_arg,
                                                                                 self.ai_provider_key,
                                                                                 self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_PDF_CONTENT_SUMMARY,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_pdf_content_summary(df_arg, res_sf_arg, cb_arg,
                                                                                    self.ai_provider_key,
                                                                                    self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_LITERATURE_REVIEW_SUMMARY,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_literature_review_summary(df_arg, res_sf_arg, cb_arg,
                                                                                          self.ai_provider_key,
                                                                                          self.ai_model_api_name),
                 "html_section", False, True),
                (STEP_AI_DISCUSSION,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_discussion(df_arg, res_sf_arg, cb_arg,
                                                                           self.ai_provider_key,
                                                                           self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_LIMITATIONS,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_limitations(df_arg, res_sf_arg, cb_arg,
                                                                            self.ai_provider_key,
                                                                            self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_INTRODUCTION,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_introduction(df_arg, res_sf_arg, cb_arg,
                                                                             self.ai_provider_key,
                                                                             self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_AI_ABSTRACT, lambda df_arg, cb_arg, res_sf_arg: generate_ai_abstract(df_arg, res_sf_arg, cb_arg,
                                                                                           self.ai_provider_key,
                                                                                           self.ai_model_api_name),
                 "html_section", True, True),
                (STEP_AI_CONCLUSION,
                 lambda df_arg, cb_arg, res_sf_arg: generate_ai_conclusion(df_arg, res_sf_arg, cb_arg,
                                                                           self.ai_provider_key,
                                                                           self.ai_model_api_name), "html_section",
                 True, True),
                (STEP_ASSEMBLE_REPORT_HTML,
                 lambda df_arg, cb_arg, res_sf_arg: assemble_report_html(self._bibliographic_data_df, res_sf_arg,
                                                                         cb_arg, theme_name=self.current_theme_name),
                 "full_html_report", True, True),
            ]
        }

    def _emit_log(self, message):
        self.log_message_ready.emit(message)
        if hasattr(self, '_current_step_logs_buffer'):
            self._current_step_logs_buffer.append(f"[{QDateTime.currentDateTime().toString('hh:mm:ss.zzz')}] {message}")
        QApplication.processEvents()

    def _emit_data_summary(self, df_full_data_for_summary, step_name_for_summary):  # ... (same as before)
        self.data_summary_ready.emit(df_full_data_for_summary, step_name_for_summary)


    @pyqtSlot()
    def run_analysis(self):
        self._is_running = True
        self._is_paused = False
        self._results_so_far = {}  # Initialize results for this run
        try:
            current_phase_for_load = PHASE_RESEARCH
            self.step_started.emit(current_phase_for_load, STEP_LOAD_DATA)

            load_config = next((s for s in self.full_review_steps_config[PHASE_RESEARCH] if s[0] == STEP_LOAD_DATA),
                               None)
            if not load_config:
                self._emit_log(f"Config for {STEP_LOAD_DATA} not found.")
                self.analysis_finished.emit(False, "Internal config error.")
                self._is_running = False
                return

            load_func = load_config[1]  # This is load_data_from_source_for_widget

            _load_result = load_func(
                source_type=self.data_source_type,
                file_path=self.data_file_path,
                collection_name=self.zotero_collection_name,
                progress_callback=self._emit_log
            )

            if not self._is_running:
                raise InterruptedException("Stopped during data loading")

                # ── Normalize loader output robustly ────────────────────────────────
                # Accept: pd.DataFrame | (pd.DataFrame, extras_dict) | dict payloads
                df_loaded = None
                extras = {}

                try:
                    import pandas as _pd  # local import is safe; already imported as pd at top
                except Exception:
                    _pd = pd  # fallback to existing alias

                result_type = type(_load_result).__name__
                self._emit_log(f"DEBUG: Loader returned type '{result_type}'.")

                if isinstance(_load_result, _pd.DataFrame):
                    df_loaded = _load_result

                elif isinstance(_load_result, tuple):
                    # Expect (df, extras) or (df,)
                    if len(_load_result) >= 1 and isinstance(_load_result[0], _pd.DataFrame):
                        df_loaded = _load_result[0]
                    if len(_load_result) >= 2 and isinstance(_load_result[1], dict):
                        extras.update(_load_result[1])

                elif isinstance(_load_result, dict):
                    # Try common keys where a DataFrame may live
                    candidate_keys = ("df", "dataframe", "raw_df_full", "raw_df", "data")
                    for k in candidate_keys:
                        v = _load_result.get(k)
                        if isinstance(v, _pd.DataFrame):
                            df_loaded = v
                            self._emit_log(f"DEBUG: Using DataFrame from loader payload key '{k}'.")
                            break

                    # Collect potential extras/meta if provided
                    for k in ("extras", "meta", "metadata", "info"):
                        v = _load_result.get(k)
                        if isinstance(v, dict):
                            extras.update(v)

                    # Keep whole payload for downstream reference
                    extras.setdefault("loader_payload", _load_result)

                else:
                    # Unknown return type
                    self._emit_log(f"WARNING: Unexpected loader return type '{result_type}'. Unable to normalize.")

                # Commit normalized results
                self._bibliographic_data_df = df_loaded
                if isinstance(extras, dict) and extras:
                    self._results_so_far.update(extras)

                # Final guard
                if self._bibliographic_data_df is None or getattr(self._bibliographic_data_df, "empty", True):
                    self._emit_log("Failed to load data or data is empty.")
                    self.step_completed.emit(current_phase_for_load, STEP_LOAD_DATA, False, None)
                    self.analysis_finished.emit(False, "Data loading failed.")
                    self._is_running = False
                    return

            # --- Determine dataset_title and base_output_dir for this run ---
            raw_dataset_title = "Untitled_Dataset"
            if hasattr(self, 'research_topic') and self.research_topic:
                raw_dataset_title = self.research_topic
            elif self.data_source_type == "zotero" and self.zotero_collection_name:
                raw_dataset_title = self.zotero_collection_name
            elif self.data_file_path:
                try:
                    raw_dataset_title = Path(self.data_file_path).stem
                except:
                    raw_dataset_title = "File_Dataset"

            def sanitize_for_path_local(name_str: str, max_len: int = 100) -> str:  # Local sanitize for this scope
                if not name_str: name_str = "untitled_item"
                name_str = str(name_str)
                name_str = re.sub(r'[<>:"/\\|?*\s&]', '_', name_str)
                name_str = re.sub(r'_+', '_', name_str).strip('_')
                return name_str[:max_len]

            sanitized_dataset_title_for_folder = sanitize_for_path_local(raw_dataset_title)

            # Define base output path: output_default/sanitized_collection_name/
            # User mentioned "output_default" as a desired root.
            root_output_directory = Path.cwd() / "output_default"
            this_run_base_output_dir = root_output_directory / sanitized_dataset_title_for_folder
            this_run_base_output_dir.mkdir(parents=True, exist_ok=True)
            self._emit_log(f"Base output directory for this run set to: {this_run_base_output_dir}")
            # Always operate on a safe DataFrame to avoid calling .copy()/.head() on None
            df_safe = self._bibliographic_data_df if isinstance(self._bibliographic_data_df,
                                                                pd.DataFrame) else pd.DataFrame()

            try:
                data_html_preview = df_safe.head().to_html(classes='dataframe small-table', border=0, justify='left')
            except Exception:
                data_html_preview = "<p class='placeholder-text'>[No data preview available]</p>"

            loaded_data_output = {
                "type": "raw_df_full_summary",
                "raw_df_full": df_safe.copy(),
                "data_html": data_html_preview,
                "description": f"Loaded Dataset Overview: {raw_dataset_title}",
                "full_df_shape": getattr(df_safe, "shape", (0, 0)),
                "collection_name_for_title": raw_dataset_title,
                "data_source_type_used": self.data_source_type,
                "original_zotero_collection_name": self.zotero_collection_name if self.data_source_type == "zotero" else None,
                "original_file_path": self.data_file_path if self.data_source_type == "file" else None,
                "base_output_dir": str(this_run_base_output_dir)  # Stored for other steps
            }
            self._results_so_far[STEP_LOAD_DATA] = loaded_data_output  # This makes it available
            self._emit_data_summary(df_safe.copy(), STEP_LOAD_DATA)

            self.step_completed.emit(current_phase_for_load, STEP_LOAD_DATA, True, loaded_data_output)
            self._check_pause_and_stop(current_phase_for_load, STEP_LOAD_DATA, True)

            ## MODIFICATION: Route to the correct runner based on analysis type
            if self.analysis_type == ANALYSIS_TYPE_FULL_REVIEW:
                self._run_full_review_phased()
            elif self.analysis_type == ANALYSIS_TYPE_CHRONOLOGICAL_REVIEW:
                self._run_chronological_review()
            else:
                self._run_single_analysis()

            if self._is_running:
                self.analysis_finished.emit(True, "Analysis completed.")

        except InterruptedException:
            self._emit_log("Analysis was interrupted by user.")
            self.analysis_finished.emit(False, "Analysis stopped by user.")
        except Exception as e:
            self._emit_log(f"Unexpected error in run_analysis: {str(e)}\n{traceback.format_exc()}")
            self.analysis_finished.emit(False, f"Analysis failed due to an unexpected error: {e}")
        finally:
            self._is_running = False
    def _run_single_analysis(self):  # ... (same as before)
        if not self._is_running: return
        df = self._bibliographic_data_df
        single_analysis_map = {
            ANALYSIS_TYPE_AUTHOR: (STEP_ANALYZE_AUTHORS, analyze_authors_detailed_plotly, "plotly_graph_df"),
            ANALYSIS_TYPE_KEYWORD: (STEP_ANALYZE_KEYWORDS, analyze_keywords_detailed_plotly, "plotly_graph_df"),
            ANALYSIS_TYPE_TIMELINE: (STEP_GENERATE_TIMELINE,
                                     lambda d, cb, rsf: generate_timeline_detailed_plotly(d, cb), "plotly_graph"),
            ANALYSIS_TYPE_COAUTHOR_NET: (STEP_GENERATE_COAUTHORSHIP_NET,
                                         lambda d, cb, rsf: generate_network_graph_detailed_plotly(d, "authors", cb),
                                         "plotly_graph"),
        }
        if self.analysis_type not in single_analysis_map:
            self._emit_log(f"Analysis type '{self.analysis_type}' not configured for single run.");
            self.analysis_finished.emit(False, "Unknown single analysis.");
            self._is_running = False;
            return
        step_name, func_or_lambda, output_type_hint = single_analysis_map[self.analysis_type]
        current_phase = PHASE_RESEARCH
        self.step_started.emit(current_phase, step_name)
        if not self._is_running: raise InterruptedException("Stopped before single step exec")
        output_package = self._execute_step_logic(func_or_lambda, df, output_type_hint, step_name)
        if not self._is_running and (output_package is None or output_package.get("type") != "error"):
            self.step_completed.emit(current_phase, step_name, False,
                                     {"type": "interrupted", "data": "Step interrupted."});
            raise InterruptedException("Stopped during single step exec")
        success = output_package is not None and output_package.get("type") != "error"
        self.step_completed.emit(current_phase, step_name, success,
                                 output_package if success else {"type": "error", "data": f"Failed: {step_name}"})
        if success:
            self._results_so_far[step_name] = output_package; self._check_pause_and_stop(current_phase, step_name, True)
        else:
            self._emit_log(f"Single analysis '{step_name}' failed."); self._is_running = False

    def _execute_step_logic(self, func_or_lambda, df_input, output_type_hint, step_name, phase_name=""):
        stored_output_package = None
        try:
            if hasattr(self, '_current_step_logs_buffer'):
                self._current_step_logs_buffer.clear()

            if not self._is_running:
                raise InterruptedException(f"Execution attempt for {step_name} while not running.")

            # Defensive: accept df_input as DataFrame or (DataFrame, …)
            # Defensive: accept df_input as DataFrame or (DataFrame, …)
            if isinstance(df_input, tuple):
                try:
                    df_input = df_input[0]
                except Exception:
                    df_input = None  # ensure we handle failures gracefully

            # Normalize to a safe DataFrame so downstream funcs never see None
            if not isinstance(df_input, pd.DataFrame):
                self._emit_log(f"WARNING: Step '{step_name}' received no valid DataFrame; using empty DataFrame.")
                df_input = pd.DataFrame()

            # Ensure results dict is never None to avoid: 'NoneType' object has no attribute 'get'
            safe_results = self._results_so_far or {}
            args_for_lambda = (df_input, self._emit_log, safe_results)

            def sanitize_for_path(name_str: str, max_len: int = 100) -> str:
                if not name_str: name_str = "untitled_item"
                name_str = str(name_str)
                name_str = re.sub(r'[<>:"/\\|?*\s&]', '_', name_str)
                name_str = re.sub(r'_+', '_', name_str).strip('_')
                return name_str[:max_len]

            # --- Determine images folder for this run ---
            # This base_output_dir is specific to the collection, e.g., "output_default/My_Collection"
            current_run_base_dir_str = self._results_so_far.get(STEP_LOAD_DATA, {}).get("base_output_dir")
            images_base_folder = None  # Initialize

            if current_run_base_dir_str:
                images_base_folder = Path(current_run_base_dir_str) / "images"
                try:
                    images_base_folder.mkdir(parents=True, exist_ok=True)
                    self._emit_log(f"DEBUG: Image save directory confirmed: {images_base_folder}")
                except Exception as e_mkdir:
                    self._emit_log(
                        f"ERROR: Could not create images directory {images_base_folder}: {e_mkdir}. Images will not be saved.")
                    images_base_folder = None
            else:
                self._emit_log(
                    f"WARNING: 'base_output_dir' not found in results_so_far[{STEP_LOAD_DATA}]. Cannot determine image save path for step {step_name}. Check run_analysis setup.")

            def save_plot_image(fig, current_step_name_for_file, output_pkg_ref):
                if fig is None:
                    self._emit_log(f"No figure object to save for '{current_step_name_for_file}'.")
                    if isinstance(output_pkg_ref, dict): output_pkg_ref["image_file_path"] = None
                    return

                if images_base_folder is None:
                    self._emit_log(
                        f"Skipping image save for '{current_step_name_for_file}' as images directory unavailable.")
                    if isinstance(output_pkg_ref, dict): output_pkg_ref["image_file_path"] = None
                    return

                image_path = None  # Initialize image_path to None
                try:
                    sanitized_file_name_part = sanitize_for_path(current_step_name_for_file)
                    image_filename = f"{sanitized_file_name_part}_plot.png"
                    image_path = images_base_folder / image_filename  # This is a Path object

                    self._emit_log(f"DEBUG: Target image path for {current_step_name_for_file}: {image_path}")

                    if image_path.exists():  # Check 1: Does it already exist?
                        self._emit_log(f"Using existing plot image for '{current_step_name_for_file}': {image_path}")
                        # If it exists, image_path is already correctly set to the existing file.
                    else:
                        # Try to save it
                        self._emit_log(
                            f"Attempting to save new plot image for '{current_step_name_for_file}' to {image_path}...")
                        try:
                            import kaleido  # Check for kaleido library
                            fig.write_image(str(image_path), scale=1.5)  # Convert Path object to string for this call

                            # Check 2: Critical check immediately after trying to write
                            if image_path.exists():
                                self._emit_log(f"Successfully saved plot image: {image_path}")
                            else:
                                # This is likely where your issue lies if kaleido didn't error but no file was made
                                self._emit_log(
                                    f"ERROR: Image file NOT FOUND at '{image_path}' after attempting save. Kaleido might have failed silently or an issue with the figure object prevented saving.")
                                image_path = None  # Explicitly set to None as save failed
                        except ImportError:
                            self._emit_log(
                                f"ERROR: SAVE PLOT FAILED for '{current_step_name_for_file}': 'kaleido' library not installed. Please install with 'pip install -U kaleido'. Image not saved.")
                            image_path = None  # Mark as not saved
                        except Exception as e_write:  # Catch other errors during write_image
                            self._emit_log(
                                f"ERROR: Failed to write plot image for '{current_step_name_for_file}' due to: {e_write}")
                            image_path = None  # Mark as not saved

                except Exception as e_img_setup:  # Catch errors in path creation or other setup
                    self._emit_log(
                        f"ERROR: Outer error during plot image saving setup for '{current_step_name_for_file}': {e_img_setup}")
                    image_path = None  # Ensure image_path is None on any setup error

                finally:
                    # This block ensures 'image_file_path' is always set in the output package.
                    # It will be None if image_path is None OR if image_path is set but the file doesn't actually exist.
                    final_path_to_store = str(image_path) if image_path and image_path.exists() else None
                    if isinstance(output_pkg_ref, dict):
                        output_pkg_ref["image_file_path"] = final_path_to_store
                    self._emit_log(
                        f"DEBUG: Final 'image_file_path' stored in package for step '{current_step_name_for_file}': {final_path_to_store}")
            # --- Main Step Logic ---
            # Process STEP_LOAD_DATA (primarily done in run_analysis, this is defensive)
            if output_type_hint == "raw_df_full_summary" and step_name == STEP_LOAD_DATA:
                    # or this is for a single-step "Load Data" only analysis.
                self._emit_log(
                    f"Executing STEP_LOAD_DATA within _execute_step_logic. This is unusual for full review.")

                stored_output_package = self._results_so_far.get(STEP_LOAD_DATA)  # Use what run_analysis put there.
                if not stored_output_package:  # If run_analysis somehow didn't set it (e.g. single analysis)
                    # This part would replicate the package creation from run_analysis.
                    # For simplicity, assuming run_analysis always correctly populates STEP_LOAD_DATA package.
                    self._emit_log(
                        f"CRITICAL WARNING: STEP_LOAD_DATA package not found in _results_so_far where expected for {step_name}.")

                    df_loaded_via_execute = func_or_lambda(source_type=self.data_source_type,
                                                           file_path=self.data_file_path,
                                                           collection_name=self.zotero_collection_name,
                                                           progress_callback=self._emit_log)
                    raw_dataset_title_exec = self._results_so_far.get(STEP_LOAD_DATA, {}).get(
                        "collection_name_for_title", "Default_Execute_Topic")
                    base_output_dir_exec = self._results_so_far.get(STEP_LOAD_DATA, {}).get("base_output_dir",
                                                                                            str(Path.cwd() / "output_default" / sanitize_for_path(
                                                                                                raw_dataset_title_exec)))
                    stored_output_package = {
                        "type": "raw_df_full_summary",
                        "raw_df_full": df_loaded_via_execute.copy() if df_loaded_via_execute is not None else pd.DataFrame(),
                        "description": f"Loaded Data: {raw_dataset_title_exec}",
                        "collection_name_for_title": raw_dataset_title_exec,
                        "base_output_dir": base_output_dir_exec
                        # ... plus other fields from run_analysis's loaded_data_output
                    }
                    self._results_so_far[STEP_LOAD_DATA] = stored_output_package


            elif output_type_hint == "plotly_graph_df":

                # Short-circuit: if the normalized DF is empty, emit a friendly placeholder

                _df_for_step = args_for_lambda[0]

                if not isinstance(_df_for_step, pd.DataFrame) or _df_for_step.empty:

                    stored_output_package = {

                        "type": "table_summary",

                        "raw_df_table": pd.DataFrame(),  # empty

                        "description": f"No data available for {step_name}",

                        "data_html_fragment": "<p class='placeholder-text'>[No data available to plot for this step]</p>"

                    }

                    # still let the UI know something happened

                    self.table_data_ready.emit(pd.DataFrame(), f"Data: {step_name} (Empty)")

                else:

                    df_res, fig_plotly = func_or_lambda(*args_for_lambda)

                    if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")

                    stored_output_package = {}

                    if fig_plotly is not None:
                        self.plotly_graph_ready.emit(fig_plotly, f"Chart: {step_name}")

                        stored_output_package.update({

                            "type": "plotly_html_summary",

                            "data_html_fragment": fig_plotly.to_html(full_html=False, include_plotlyjs='cdn'),

                            "description": f"Plot for {step_name}"

                        })

                        save_plot_image(fig_plotly, step_name, stored_output_package)

                    if df_res is not None and not df_res.empty:

                        self.table_data_ready.emit(df_res, f"Data: {step_name}")

                        stored_output_package["raw_df_table"] = df_res.copy()

                        if "type" not in stored_output_package:
                            stored_output_package["type"] = "table_summary"

                            stored_output_package["description"] = f"Table for {step_name}"

                    if not stored_output_package and (df_res is not None or fig_plotly is not None):
                        stored_output_package = {"type": "data_package", "description": f"Data for {step_name}"}

            ## MODIFICATION: Add a handler for the new 'period_analysis_package' output type
            elif output_type_hint == "period_analysis_package":
                # This step returns a dictionary of analyses for different time periods.
                period_results = func_or_lambda(*args_for_lambda)
                if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")

                # The package itself is the result. We can also emit live data for each period.
                for period_name, period_data in period_results.items():
                    self._emit_log(f"Generated analysis for period: {period_name}")
                    if period_data.get("figure"):
                        self.live_plotly_ready.emit(period_data["figure"], f"{step_name}: {period_name} Chart")
                    if period_data.get("table") is not None:
                        self.live_table_ready.emit(period_data["table"], f"{step_name}: {period_name} Data")

                stored_output_package = {
                    "type": "period_analysis_package",
                    "data": period_results,
                    "description": "Analysis of Key Publication Periods"
                }

            elif output_type_hint == "plotly_graph":  # Handles STEP_GENERATE_TIMELINE
                fig_plotly = func_or_lambda(*args_for_lambda)
                if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")
                stored_output_package = {  # Initialize with defaults
                    "type": "plotly_html_summary", "description": f"Plot for {step_name}",
                    "data_html_fragment": None, "image_file_path": None,
                    "_data_summary_text": "No trend data to summarize."  # Default for timeline
                }
                if fig_plotly is not None:
                    self.plotly_graph_ready.emit(fig_plotly, f"Chart: {step_name}")
                    stored_output_package["data_html_fragment"] = fig_plotly.to_html(full_html=False,
                                                                                     include_plotlyjs='cdn')
                    save_plot_image(fig_plotly, step_name,
                                    stored_output_package)  # Saves image and sets "image_file_path"

                    if step_name == STEP_GENERATE_TIMELINE:
                        if df_input is not None and not df_input.empty and 'year' in df_input.columns:
                            try:
                                # Ensure analyze_publication_trends & describe_trend_data are imported or accessible
                                pub_trends_series = analyze_publication_trends(df_input)  # from data_processing
                                if not pub_trends_series.empty:
                                    detailed_summary = describe_trend_data(
                                        pub_trends_series)  # from data_processing
                                    stored_output_package["_data_summary_text"] = detailed_summary
                                    self._emit_log(
                                        f"DEBUG: Generated detailed textual summary for {step_name} (len: {len(detailed_summary)}).")
                                else:
                                    self._emit_log(
                                        f"DEBUG: Publication trends series is empty for {step_name}. Using default summary.")
                            except NameError as ne:  # If functions aren't imported
                                self._emit_log(
                                    f"ERROR: Could not generate detailed timeline summary: {ne}. Functions might not be imported into widget worker.")
                            except Exception as e_sum:
                                self._emit_log(f"ERROR: Exception generating detailed timeline summary: {e_sum}")
                        else:
                            self._emit_log(
                                f"DEBUG: DataFrame for {step_name} is empty or missing 'year' column. Cannot generate detailed summary.")
                else:
                    self._emit_log(f"Figure generation failed for {step_name}. No plot to display or summarize.")


            elif output_type_hint == "plotly_graph_dict":  # Handles network graphs
                graph_package = func_or_lambda(*args_for_lambda)
                fig_plotly = None
                aux_data_to_store_in_package = {}
                stored_output_package = {}

                if isinstance(graph_package, dict):
                    fig_plotly = graph_package.get("figure")
                    aux_data_to_store_in_package = {k: v for k, v in graph_package.items() if k != "figure"}
                    # Storing cluster info for keyword co-occurrence network
                    if "keyword_clusters_data" in graph_package and step_name == STEP_GENERATE_KEYWORD_COOCCURRENCE_NET:
                        # This structure is directly put into results_so_far for the step,
                        # but also useful to have it directly linked if ai_services expects it this way.
                        self._results_so_far[f"{STEP_GENERATE_KEYWORD_COOCCURRENCE_NET}_clusters_info"] = \
                        graph_package["keyword_clusters_data"]
                        stored_output_package["keyword_clusters_data"] = graph_package[
                            "keyword_clusters_data"]  # Add to current step package too
                        self._emit_log(f"Keyword cluster information stored for {step_name}.")

                if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")

                if fig_plotly is not None:
                    self.plotly_graph_ready.emit(fig_plotly, f"Chart: {step_name}")
                    stored_output_package.update({
                        "type": "plotly_html_summary",
                        "data_html_fragment": fig_plotly.to_html(full_html=False, include_plotlyjs='cdn'),
                        "description": f"Plot for {step_name}"
                    })
                    save_plot_image(fig_plotly, step_name, stored_output_package)

                if aux_data_to_store_in_package:
                    if stored_output_package:
                        stored_output_package.update(aux_data_to_store_in_package)
                    else:
                        stored_output_package = {
                            "type": f"{sanitize_for_path(step_name)}_data_only",
                            "data": aux_data_to_store_in_package,
                            "description": f"Auxiliary Data for {step_name}"
                        }
                        # Ensure a type if it's only aux_data
                        if "type" not in stored_output_package and aux_data_to_store_in_package:
                            stored_output_package["type"] = "aux_data_summary"

            elif output_type_hint in ["keyword_list", "author_keywords_map", "keyword_cluster_search_terms_map"]:
                if args_for_lambda and len(args_for_lambda) >= 3:
                    df_arg, log_cb, res_sf_arg = args_for_lambda
                    res_sf_arg = res_sf_arg or {}
                    args_for_lambda = (df_arg, log_cb, res_sf_arg)
                result_package_data = func_or_lambda(*args_for_lambda)
                if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")
                if result_package_data and result_package_data.get("type") not in ("error", None):
                    desc = result_package_data.get("description", step_name)
                    data_content = result_package_data.get("data", {})
                    html_block_content = ""
                    if isinstance(data_content, dict):
                        html_block_content = "<pre>" + html.escape(json.dumps(data_content, indent=2)) + "</pre>"
                    elif isinstance(data_content, list):
                        html_block_content = "<ul>" + "".join(
                            f"<li>{html.escape(str(i))}</li>" for i in data_content) + "</ul>"
                    else:
                        html_block_content = f"<p>{html.escape(str(data_content))}</p>"
                    response_html_for_package = f"<h3>{html.escape(desc)}</h3>{html_block_content}"
                    stored_output_package = {
                        "type": "html_section",
                        "data": {"response_html": response_html_for_package,
                                 "prompt_sent": result_package_data.get("prompt_sent", "N/A for utility function")},
                        "description": desc
                    }
                    self.html_section_ready.emit(stored_output_package['data'], desc)
                elif result_package_data and result_package_data.get("type") == "error":
                    stored_output_package = result_package_data
                    self.html_section_ready.emit(
                        result_package_data.get("data", {"error_message": "Unknown error from utility."}),
                        f"Error: {result_package_data.get('description', step_name)}"
                    )
                else:
                    stored_output_package = {"type": "generic_data_dict", "data": result_package_data,
                                             "description": f"Raw Data from {step_name}"}

            elif output_type_hint == "html_section":
                ai_output_dict = func_or_lambda(*args_for_lambda)
                if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")
                stored_output_package = {
                    "type": "html_section", "data": ai_output_dict,
                    "description": ai_output_dict.get("description", f"AI Analysis for {step_name}")
                }
                payload_for_monitor = (
                    ai_output_dict.get("data", ai_output_dict)  # pull inner dict if still wrapped
                    if isinstance(ai_output_dict, dict) else ai_output_dict
                )

                payload = (
                    ai_output_dict.get("data", ai_output_dict)  # unwrap if nested
                    if isinstance(ai_output_dict, dict) else ai_output_dict
                )

                # supply whichever alias is missing
                payload.setdefault("response_html",
                                   payload.get("initial_draft_html_response", ""))
                payload.setdefault("initial_draft_html_response",
                                   payload.get("response_html", ""))
                payload.setdefault("prompt_sent",
                                   payload.get("initial_draft_prompt_sent", ""))
                payload.setdefault("initial_draft_prompt_sent",
                                   payload.get("prompt_sent", ""))

                self.html_section_ready.emit(payload,
                                             stored_output_package["description"])
            elif output_type_hint == "full_html_report":
                html_report_string = func_or_lambda(*args_for_lambda)
                if not self._is_running: raise InterruptedException(f"Interrupted during {step_name}")
                stored_output_package = {
                    "type": "full_html_report", "data": html_report_string,
                    "description": "Full Assembled Bibliographic Report"
                }
                self.html_section_ready.emit({"response_html": html_report_string}, STEP_ASSEMBLE_REPORT_HTML)
            else:  # Fallback for unhandled types
                # Check if it's an unhandled plot type or truly unknown
                if not stored_output_package:  # If none of the above specific handlers created it
                    self._emit_log(
                        f"Output type hint '{output_type_hint}' for {step_name} not specifically handled.")
                    raw_result = func_or_lambda(*args_for_lambda)
                    stored_output_package = {"type": "unknown_output", "data": raw_result,
                                             "description": f"Raw output from {step_name}"}

            # Add input summary and step-specific logs
            if stored_output_package and isinstance(stored_output_package, dict):
                if hasattr(self, '_current_step_logs_buffer'):
                    stored_output_package["step_specific_logs"] = list(self._current_step_logs_buffer)
                # df_input for STEP_LOAD_DATA is not relevant in this function as initial load happens in run_analysis
                if step_name != STEP_LOAD_DATA and df_input is not None:
                    input_desc = f"Input: DataFrame with {df_input.shape[0]} rows, {df_input.shape[1]} columns."
                    stored_output_package["input_summary_text"] = input_desc
                    if not df_input.empty:
                        try:
                            stored_output_package["input_data_sample_html"] = df_input.head(3).to_html(
                                classes='dataframe small-table', border=0, justify='left', max_cols=7)
                        except Exception as e_sample:
                            self._emit_log(
                                f"Warning: Could not generate input data sample for {step_name}: {e_sample}")

            if stored_output_package is None:
                stored_output_package = {"type": "no_output", "data": f"No structured output for {step_name}.",
                                         "description": f"{step_name} (No Output)"}
            # --- End of copied section ---

        except InterruptedException as ie:
            self._emit_log(f"Step '{step_name}' interrupted: {ie}")
            return {"type": "interrupted", "data": {"error_message": f"Interrupted: {step_name}"},
                    "description": f"Interrupted: {step_name}"}
        except Exception as e_step:
            self._emit_log(f"Critical error in logic of step '{step_name}': {str(e_step)}")
            logging.error(f"Critical error in worker step {step_name}: {e_step}", exc_info=True)
            return {"type": "error",
                    "data": {"error_message": f"Unhandled Exception in {step_name}: {str(e_step)}",
                             "traceback": traceback.format_exc()},
                    "description": f"Critical Error in {step_name}"}

        return stored_output_package

    ## MODIFICATION: Create a new runner method for the chronological review
    def _run_chronological_review(self):
        df_current = self._bibliographic_data_df
        for phase_name, steps_in_phase in self.chronological_review_steps_config.items():
            if not self._is_running: raise InterruptedException(f"Stopped before {phase_name}")
            self._emit_log(f"--- Entering {phase_name} ---")
            for step_config in steps_in_phase:
                if not self._is_running: raise InterruptedException(f"Stopped during {phase_name}")
                step_name, func_or_lambda, output_type_hint, is_major_step, stores_output = step_config
                if step_name == STEP_LOAD_DATA: continue
                self.step_started.emit(phase_name, step_name)
                if not self._is_running: raise InterruptedException(f"Stopped before step {step_name}")
                output_package = self._execute_step_logic(func_or_lambda, df_current, output_type_hint, step_name,
                                                          phase_name=phase_name)
                if not self._is_running and (output_package is None or output_package.get("type") != "error"):
                    self.step_completed.emit(phase_name, step_name, False,
                                             {"type": "interrupted", "data": "Step interrupted."});
                    raise InterruptedException(f"Stopped during step {step_name}")
                success = output_package is not None and output_package.get("type") != "error"
                if success and stores_output: self._results_so_far[step_name] = output_package
                self.step_completed.emit(phase_name, step_name, success,
                                         output_package if success else {"type": "error",
                                                                         "data": f"Error in {step_name}"})
                if not success and is_major_step:
                    self._emit_log(f"Major step '{step_name}' in {phase_name} failed. Aborting.");
                    self._is_running = False
                    raise Exception(f"Major step {step_name} failed.")
                if success and is_major_step: self._check_pause_and_stop(phase_name, step_name, is_major_step)


    def _run_full_review_phased(self):  # ... (same as before, uses self.full_review_steps_config)
        df_current = self._bibliographic_data_df
        for phase_name, steps_in_phase in self.full_review_steps_config.items():
            if not self._is_running: raise InterruptedException(f"Stopped before {phase_name}")
            self._emit_log(f"--- Entering {phase_name} ---")
            for step_config in steps_in_phase:
                if not self._is_running: raise InterruptedException(f"Stopped during {phase_name}")
                step_name, func_or_lambda, output_type_hint, is_major_step, stores_output = step_config
                if step_name == STEP_LOAD_DATA: continue
                self.step_started.emit(phase_name, step_name)
                if not self._is_running: raise InterruptedException(f"Stopped before step {step_name}")
                output_package = self._execute_step_logic(func_or_lambda, df_current, output_type_hint, step_name,
                                                          phase_name=phase_name)
                if not self._is_running and (output_package is None or output_package.get("type") != "error"):
                    self.step_completed.emit(phase_name, step_name, False,
                                             {"type": "interrupted", "data": "Step interrupted."});
                    raise InterruptedException(f"Stopped during step {step_name}")
                success = output_package is not None and output_package.get("type") != "error"
                if success and stores_output: self._results_so_far[step_name] = output_package
                self.step_completed.emit(phase_name, step_name, success,
                                         output_package if success else {"type": "error",
                                                                         "data": f"Error in {step_name}"})
                if not success and is_major_step:
                    self._emit_log(f"Major step '{step_name}' in {phase_name} failed. Aborting.");
                    self._is_running = False
                    raise Exception(f"Major step {step_name} failed.")
                if success and is_major_step: self._check_pause_and_stop(phase_name, step_name, is_major_step)



    def _check_pause_and_stop(self, phase_name, step_name, is_major_step_for_manual_pause):  # ... (same as before)
        if not self._is_running: raise InterruptedException(f"Pre-check stop for {step_name}")
        self._mutex.lock()
        try:
            if not self._is_running: self._mutex.unlock(); raise InterruptedException(
                f"Pre-check stop (locked) for {step_name}")
            if (self.is_manual_mode and is_major_step_for_manual_pause) or self._is_paused:
                actual_pause_step_name = step_name if self.is_manual_mode and is_major_step_for_manual_pause else f"User Pause during {step_name}"
                self.paused_for_manual_step.emit(phase_name, actual_pause_step_name)
                self._is_paused = True
                while self._is_paused and self._is_running: self._pause_condition.wait(self._mutex)
            if not self._is_running: self._mutex.unlock(); raise InterruptedException(f"Post-wait stop for {step_name}")
        finally:
            self._mutex.unlock()

    def resume_manual_step(self):  # ... (same as before)
        self._mutex.lock();
        self._is_paused = False;
        self._pause_condition.wakeAll();
        self._mutex.unlock();
        self._emit_log("Resuming...")

    def toggle_pause_resume(self):  # ... (same as before)
        self._mutex.lock()
        if not self._is_running: self._is_paused = False; self._mutex.unlock(); return self._is_paused
        self._is_paused = not self._is_paused
        if not self._is_paused:
            self._pause_condition.wakeAll(); self._emit_log("Resumed by user.")
        else:
            self._emit_log("Paused by user.")
        self._mutex.unlock();
        return self._is_paused

    def request_stop(self):  # ... (same as before)
        self._emit_log("Stop requested.");
        self._mutex.lock()
        self._is_running = False;
        self._is_paused = False;
        self._pause_condition.wakeAll();
        self._mutex.unlock()


class InterruptedException(Exception): pass

#
class BibliographicAnalysisWidget(QWidget):

    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_research_topic = "cyber attribution"

        self.setObjectName("bibliographicAnalysisWidget")
        self.bibliographic_data_df = None
        self.current_analysis_type = None
        self.worker_thread = None
        self.analysis_worker = None  # This will be an instance of BibliographicAnalysisWorker
        self.is_manual_mode = False
        self.current_theme_name = "dark"  # Default theme
        default_provider = list(AI_MODELS_CONFIG.keys())[0] if AI_MODELS_CONFIG else "OpenAI"
        default_models_for_provider = AI_MODELS_CONFIG.get(default_provider, {})
        default_model_display_name = list(default_models_for_provider.keys())[
            0] if default_models_for_provider else "Default Model"
        self.current_ai_provider_key = default_provider
        self.current_ai_model_api_name = default_models_for_provider.get(default_model_display_name,
                                                                         "default-api-model-name")
        self.completed_step_outputs = {}  # Stores the final output package for each completed step
        self.executed_steps_chronological = []  # Keeps track of step completion order
        self.current_display_step_tuple = None  # (phase_name, step_name) of the step being viewed
        self.current_output_views = []  # List of (view_type_str, view_content, view_title_suffix_str) for the current_display_step_tuple
        self.current_output_view_index = 0  # Index for multi-page views

        self._current_step_live_buffer: list[tuple[str, object, str]] = []
        self._current_phase_running: str | None = None
        self._current_step_running: str | None = None

        self._setup_ui()
        self._connect_signals()
        self._update_analysis_controls_state()
        self.apply_theme(self.current_theme_name)
        self._populate_ai_provider_combo()
        self._harden_ai_model_selection()
    def _harden_ai_model_selection(self):
        """
        Ensure provider/model combos and internal fields are always valid.
        Falls back to DEFAULT_AI_PROVIDER / DEFAULT_AI_MODEL env vars or safe OpenAI mini.
        """
        import os, logging
        global AI_MODELS_CONFIG
        try:
            default_provider = os.getenv("DEFAULT_AI_PROVIDER", "OpenAI")
            default_model_api = os.getenv("DEFAULT_AI_MODEL", "gpt-4o-mini")

            # Seed minimal config if necessary
            if not isinstance(AI_MODELS_CONFIG, dict) or not AI_MODELS_CONFIG:
                AI_MODELS_CONFIG = {
                    "OpenAI": {"GPT-4o mini": "gpt-4o-mini", "GPT-4o": "gpt-4o"}
                }
                logging.warning("AI_MODELS_CONFIG was empty. Seeded minimal defaults.")

            prov_combo = getattr(self, "ai_provider_combo", None)
            model_combo = getattr(self, "ai_model_combo", None)

            # Populate provider combo if empty
            if prov_combo is not None and prov_combo.count() == 0:
                for prov in AI_MODELS_CONFIG.keys():
                    prov_combo.addItem(prov)

            # Choose provider
            chosen_provider = (
                default_provider if default_provider in AI_MODELS_CONFIG else
                (prov_combo.currentText() if prov_combo and prov_combo.count() else next(iter(AI_MODELS_CONFIG.keys())))
            )

            # Populate models for chosen provider if empty
            provider_models = AI_MODELS_CONFIG.get(chosen_provider, {})
            if model_combo is not None and model_combo.count() == 0:
                for disp in provider_models.keys():
                    model_combo.addItem(disp)

            # Choose model api
            chosen_disp = model_combo.currentText() if (model_combo and model_combo.currentText()) else None
            if chosen_disp and chosen_disp in provider_models:
                chosen_model_api = provider_models[chosen_disp]
            else:
                if provider_models:
                    chosen_disp = next(iter(provider_models.keys()))
                    if model_combo is not None:
                        model_combo.setCurrentText(chosen_disp)
                    chosen_model_api = provider_models[chosen_disp]
                else:
                    chosen_model_api = default_model_api

            # Set fields
            self.current_ai_provider_key = chosen_provider
            self.current_ai_model_api_name = chosen_model_api
            logging.info(
                f"AI config seeded with defaults: Provider='{self.current_ai_provider_key}', Model API='{self.current_ai_model_api_name}'"
            )

            # Inform downstream logic (best-effort)
            try:
                self._on_ai_model_changed()
            except Exception:
                pass

        except Exception as e:
            try:
                logging.error(f"_harden_ai_model_selection error: {e}")
            except Exception:
                pass

    def _setup_ui(self):
        main_layout = QHBoxLayout(self)
        left_panel = QFrame();
        left_panel.setObjectName("analysisControlsPanel");
        left_panel.setFrameShape(QFrame.Shape.StyledPanel)
        left_panel.setMinimumWidth(400);
        left_panel.setMaximumWidth(500)
        left_layout = QVBoxLayout(left_panel)
        data_source_group = QGroupBox("1. Data Source");
        data_source_layout = QVBoxLayout(data_source_group)
        self.zotero_collection_input = QLineEdit();
        self.zotero_collection_input.setPlaceholderText("Zotero: Collection name (blank for all)")
        self.zotero_collection_input.setText("0.13_cyber_attribution_corpus_records_total_included")

        self.topic_input_label = QLabel("Research Topic/Focus:")
        data_source_layout.addWidget(self.topic_input_label)

        self.topic_input_field = QLineEdit()
        self.topic_input_field.setPlaceholderText("Enter the primary research topic (e.g., 'climate change impact')")
        self.topic_input_field.setText("cyber attribution")
        data_source_layout.addWidget(self.topic_input_field)

        self.load_from_zotero_btn = QPushButton(QIcon.fromTheme("network-server"), "Load from Zotero")
        self.load_file_btn = QPushButton(QIcon.fromTheme("document-open"), "Load from File (.csv, .xlsx)")
        self.data_status_label = QLabel("No data source indicated.");
        self.data_status_label.setObjectName("dataStatusLabel")
        data_source_layout.addWidget(self.zotero_collection_input);
        data_source_layout.addWidget(self.load_from_zotero_btn)
        data_source_layout.addWidget(self.load_file_btn);
        data_source_layout.addWidget(self.data_status_label)
        left_layout.addWidget(data_source_group)
        ai_config_group = QGroupBox("2. AI Configuration");
        ai_config_layout = QVBoxLayout(ai_config_group)
        provider_model_layout = QHBoxLayout();
        self.ai_provider_combo = QComboBox();
        provider_model_layout.addWidget(QLabel("Provider:"))
        provider_model_layout.addWidget(self.ai_provider_combo, 1);
        self.ai_model_combo = QComboBox();
        provider_model_layout.addWidget(QLabel("Model:"))
        provider_model_layout.addWidget(self.ai_model_combo, 2);
        ai_config_layout.addLayout(provider_model_layout)
        left_layout.addWidget(ai_config_group)
        analysis_type_group = QGroupBox("3. Analysis Type");
        analysis_type_layout = QVBoxLayout(analysis_type_group)
        self.analysis_type_combo = QComboBox()
        ## MODIFICATION: Add the new analysis type to the dropdown menu
        self.analysis_type_combo.addItems([
            ANALYSIS_TYPE_CHRONOLOGICAL_REVIEW,
            ANALYSIS_TYPE_FULL_REVIEW,
            ANALYSIS_TYPE_AUTHOR,
            ANALYSIS_TYPE_KEYWORD,
            ANALYSIS_TYPE_TIMELINE,
            ANALYSIS_TYPE_COAUTHOR_NET
        ])
        analysis_type_layout.addWidget(self.analysis_type_combo);
        left_layout.addWidget(analysis_type_group)
        mode_run_group = QGroupBox("4. Execution & Navigation");
        mode_run_layout = QVBoxLayout(mode_run_group)
        mode_hbox = QHBoxLayout();
        self.auto_mode_radio = QRadioButton("Automatic");
        self.manual_mode_radio = QRadioButton("Manual")
        self.auto_mode_radio.setChecked(True);
        self.mode_button_group = QButtonGroup(self);
        self.mode_button_group.addButton(self.auto_mode_radio);
        self.mode_button_group.addButton(self.manual_mode_radio)
        mode_hbox.addWidget(self.auto_mode_radio);
        mode_hbox.addWidget(self.manual_mode_radio);
        mode_run_layout.addLayout(mode_hbox)
        # Batch mode toggle (store-only pass; batch read will be triggered on finish)
        self.batch_mode_checkbox = QCheckBox("Batch mode (prepare & queue only)")
        self.batch_mode_checkbox.setToolTip(
            "If ON, AI calls are written to an OpenAI Batch input file (store-only). "
            "When the batch completes, the app re-runs in READ mode to ingest results."
        )
        mode_run_layout.addWidget(self.batch_mode_checkbox)

        self.run_analysis_btn = QPushButton(QIcon.fromTheme("media-playback-start"), "Run Analysis");
        self.run_analysis_btn.setObjectName("runAnalysisBtn");
        self.run_analysis_btn.setEnabled(False)
        mode_run_layout.addWidget(self.run_analysis_btn)
        nav_controls_hbox = QHBoxLayout();
        self.prev_step_btn = QPushButton(QIcon.fromTheme("go-previous"), "Prev Output");
        self.pause_resume_btn = QPushButton(QIcon.fromTheme("media-playback-pause"), "Pause")
        self.pause_resume_btn.setCheckable(True);
        self.next_step_btn = QPushButton(QIcon.fromTheme("go-next"), "Next Output");
        self.stop_analysis_btn = QPushButton(QIcon.fromTheme("process-stop"), "Stop")
        nav_controls_hbox.addWidget(self.prev_step_btn);
        nav_controls_hbox.addWidget(self.pause_resume_btn);
        nav_controls_hbox.addWidget(self.next_step_btn);
        nav_controls_hbox.addWidget(self.stop_analysis_btn)
        mode_run_layout.addLayout(nav_controls_hbox);
        left_layout.addWidget(mode_run_group)
        steps_group = QGroupBox("Analysis Progress");
        steps_layout = QVBoxLayout(steps_group)
        self.steps_progress_list = QListWidget();
        self.steps_progress_list.setFixedHeight(200)
        self.steps_progress_list.itemClicked.connect(self._on_step_list_item_clicked)
        steps_layout.addWidget(self.steps_progress_list);
        left_layout.addWidget(steps_group);

        report_actions_group = QGroupBox("5. Report Actions")
        report_actions_layout = QVBoxLayout(report_actions_group)

        self.export_html_btn = QPushButton(QIcon.fromTheme("document-save"), "Export Report as HTML")
        self.export_html_btn.setObjectName("exportHtmlBtn")
        self.export_html_btn.setEnabled(False)
        report_actions_layout.addWidget(self.export_html_btn)

        self.export_pdf_btn = QPushButton(QIcon.fromTheme("x-office-document-pdf"), "Export Report as PDF")
        self.export_pdf_btn.setObjectName("exportPdfBtn")
        self.export_pdf_btn.setEnabled(False)
        if not WEBENGINE_AVAILABLE:
            self.export_pdf_btn.setToolTip("PDF export requires PyQtWebEngine. Please install it.")
        report_actions_layout.addWidget(self.export_pdf_btn)

        self.export_docx_btn = QPushButton(QIcon.fromTheme("x-office-document"), "Export Report as DOCX")
        self.export_docx_btn.setObjectName("exportDocxBtn")
        self.export_docx_btn.setEnabled(False)
        report_actions_layout.addWidget(self.export_docx_btn)

        self.save_log_btn = QPushButton(QIcon.fromTheme("text-x-generic"), "Save Full Log")
        self.save_log_btn.setObjectName("saveLogBtn")
        report_actions_layout.addWidget(self.save_log_btn)

        left_layout.addWidget(report_actions_group)
        left_layout.addStretch();
        main_layout.addWidget(left_panel)
        right_panel = QFrame();
        right_panel.setObjectName("analysisMonitorPanel");
        right_panel.setFrameShape(QFrame.Shape.StyledPanel)
        right_layout = QVBoxLayout(right_panel);
        self.monitor_title_label = QLabel("Analysis Monitor");
        font = self.monitor_title_label.font();
        font.setPointSize(font.pointSize() + 2);
        font.setBold(True)
        self.monitor_title_label.setFont(font);
        self.monitor_title_label.setAlignment(Qt.AlignmentFlag.AlignCenter);
        right_layout.addWidget(self.monitor_title_label)

        self.monitor_nav_layout = QHBoxLayout()
        self.monitor_prev_page_btn = QPushButton(QIcon.fromTheme("go-previous"), "Previous Part")
        self.monitor_next_page_btn = QPushButton(QIcon.fromTheme("go-next"), "Next Part")
        self.monitor_nav_layout.addStretch(1)
        self.monitor_nav_layout.addWidget(self.monitor_prev_page_btn)
        self.monitor_nav_layout.addWidget(self.monitor_next_page_btn)
        self.monitor_nav_layout.addStretch(1)
        right_layout.addLayout(self.monitor_nav_layout)

        self.monitor_stack = QStackedWidget()
        self.monitor_placeholder_widget = QLabel("Indicate data source and run analysis, or select a completed step.");
        self.monitor_placeholder_widget.setAlignment(Qt.AlignmentFlag.AlignCenter);
        self.monitor_placeholder_widget.setStyleSheet("font-style: italic; color: grey;")
        self.monitor_stack.addWidget(self.monitor_placeholder_widget)
        self.log_monitor_browser = QTextBrowser();
        self.log_monitor_browser.setObjectName("logMonitorBrowser");
        self.monitor_stack.addWidget(self.log_monitor_browser)
        self.plotly_monitor_view = QWebEngineView()
        if WEBENGINE_AVAILABLE and self.plotly_monitor_view.settings():
            s = self.plotly_monitor_view.settings()
            s.setAttribute(QWebEngineSettings.WebAttribute.ScrollAnimatorEnabled, True)
            s.setAttribute(QWebEngineSettings.WebAttribute.JavascriptEnabled, True)
            s.setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True)
            s.setAttribute(QWebEngineSettings.WebAttribute.AllowRunningInsecureContent, True)
        self.monitor_stack.addWidget(self.plotly_monitor_view)
        self.table_monitor_widget = QTableWidget();
        self.table_monitor_widget.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers);
        self.table_monitor_widget.setAlternatingRowColors(True);
        self.table_monitor_widget.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive);
        self.table_monitor_widget.verticalHeader().setVisible(False);
        self.monitor_stack.addWidget(self.table_monitor_widget)
        self.html_content_browser = QTextBrowser()
        self.html_content_browser.setOpenExternalLinks(True)
        self.monitor_stack.addWidget(self.html_content_browser)

        right_layout.addWidget(self.monitor_stack, 1);
        main_layout.addWidget(right_panel, 1)

    def _html_with_common_css(self, raw_fragment: str) -> str:
        dark_flag = "1" if self.current_theme_name == "dark" else "0"
        return (f"<html data-dark='{dark_flag}'><head>{COMMON_MONITOR_CSS}</head>"
                f"<body>{raw_fragment}</body></html>")

    def _connect_signals(self):
        self.load_from_zotero_btn.clicked.connect(lambda: self._initiate_data_load("zotero"))
        self.load_file_btn.clicked.connect(lambda: self._initiate_data_load("file"))
        self.run_analysis_btn.clicked.connect(self._start_analysis_process)
        self.analysis_type_combo.currentTextChanged.connect(self._on_analysis_type_changed)
        self.manual_mode_radio.toggled.connect(self._on_mode_changed)
        self.ai_provider_combo.currentTextChanged.connect(self._on_ai_provider_changed)
        self.ai_model_combo.currentTextChanged.connect(self._on_ai_model_changed)
        self.pause_resume_btn.clicked.connect(self._on_pause_resume_clicked)
        self.prev_step_btn.clicked.connect(self._on_previous_step_output_clicked)
        self.next_step_btn.clicked.connect(self._on_next_step_output_clicked)
        self.stop_analysis_btn.clicked.connect(self._on_stop_analysis_clicked)
        self.monitor_prev_page_btn.clicked.connect(self._on_monitor_prev_page)
        self.monitor_next_page_btn.clicked.connect(self._on_monitor_next_page)
        try:
            self.export_html_btn.clicked.disconnect()
        except Exception:
            pass
        self.export_html_btn.clicked.connect(self._export_report_as_html)

        try:
            self.export_pdf_btn.clicked.disconnect()
        except Exception:
            pass
        self.export_pdf_btn.clicked.connect(self._export_report_as_pdf)

        try:
            self.export_docx_btn.clicked.disconnect()
        except Exception:
            pass
        self.export_docx_btn.clicked.connect(self._export_report_as_docx)
        self.save_log_btn.clicked.connect(self._save_log_to_file)

    def _start_analysis_process(self, store_only: bool | None = None, read: bool | None = None):
        self.current_research_topic = self.topic_input_field.text().strip() or "Default Research Topic"
        # ---- Batch wiring for call_models via environment flags ----
        import os
        os.environ["BATCH_FUNCTION"] = self.current_research_topic
        os.environ["BATCH_COLLECTION"] = self.current_research_topic

        if getattr(self, "batch_mode_checkbox", None) and self.batch_mode_checkbox.isChecked():
            # store-only (write JSONL, no live calls)
            os.environ["BATCH_MODE_STORE_ONLY"] = "1"
            os.environ["BATCH_MODE_READ"] = "0"
        else:
            # live (normal, immediate responses)
            os.environ["BATCH_MODE_STORE_ONLY"] = "0"
            os.environ["BATCH_MODE_READ"] = "0"

        if self.worker_thread and self.worker_thread.isRunning(): QMessageBox.warning(self, "Analysis In Progress",
                                                                                      "An analysis is already running."); return
        if not hasattr(self, 'current_data_source_type') or self.current_data_source_type is None: QMessageBox.warning(
            self, "Data Source Required", "Please indicate a data source first."); return

        self._clear_monitor_and_steps_list()
        self.completed_step_outputs.clear()
        self.executed_steps_chronological.clear()
        self.current_display_step_tuple = None
        self._current_step_live_buffer.clear()
        self._current_phase_running = None
        self._current_step_running = None

        self.log_monitor_browser.append(f"[{QDateTime.currentDateTime().toString('hh:mm:ss')}] Starting analysis...")
        self.monitor_stack.setCurrentWidget(self.monitor_placeholder_widget)

        self.current_analysis_type = self.analysis_type_combo.currentText()
        self.is_manual_mode = self.manual_mode_radio.isChecked()
        self._populate_steps_list_statically()
        self._on_ai_model_changed()
        # --- Derive batch flags ---
        if store_only is None:
            # When "Batch" is toggled on, we want to queue requests only.
            store_only = bool(getattr(self, "batch_mode_checkbox", None) and self.batch_mode_checkbox.isChecked())
        if read is None:
            # For this UX, we only queue now; reading is done by the batch runner.
            read = False
        # The BibliographicAnalysisWorker class is assumed to be defined elsewhere in the project
        # and imported correctly at the top of this file for this to work.
        self.analysis_worker = BibliographicAnalysisWorker(
            data_source_type=self.current_data_source_type,
            data_file_path=getattr(self, 'current_data_file_path', None),
            zotero_collection_name=getattr(self, 'current_zotero_collection_name', None),
            research_topic=self.current_research_topic,
            analysis_type=self.current_analysis_type,
            is_manual_mode=self.is_manual_mode,
            ai_provider_key=self.current_ai_provider_key,
            ai_model_api_name=self.current_ai_model_api_name,
            current_theme_name=self.current_theme_name,
            # NEW:
            store_only=store_only,
            read=read,
        )

        self.worker_thread = QThread(self)
        self.analysis_worker.moveToThread(self.worker_thread)

        self.analysis_worker.step_started.connect(self._handle_step_started)
        self.analysis_worker.log_message_ready.connect(self._handle_cumulative_log_message)
        self.analysis_worker.live_table_ready.connect(self._handle_live_table_data)
        self.analysis_worker.live_plotly_ready.connect(self._handle_live_plotly_data)
        self.analysis_worker.live_html_ready.connect(self._handle_live_html_data)
        self.analysis_worker.live_image_ready.connect(self._handle_live_image_data)

        self.analysis_worker.step_completed.connect(self._handle_step_completed)
        self.analysis_worker.analysis_finished.connect(self._handle_analysis_finished)
        self.analysis_worker.paused_for_manual_step.connect(self._handle_worker_paused_manual)

        self.worker_thread.started.connect(self.analysis_worker.run_analysis)
        self.worker_thread.finished.connect(self.analysis_worker.deleteLater)
        self.worker_thread.finished.connect(self.worker_thread.deleteLater)
        self.worker_thread.finished.connect(self._clear_worker_references)

        self._update_analysis_controls_state(is_running=True)
        self.worker_thread.start()

    def _generate_views_from_package(self, package: dict, step_tuple: tuple) -> list:
        views = []
        if not package or not isinstance(package, dict):
            views.append(("html_rich", f"<p>No package data for {html.escape(step_tuple[1])}.</p>", "Error"))
            return views

        step_phase, step_name = step_tuple
        output_type = package.get("type")
        output_data = package.get("data")
        output_desc = package.get("description", step_name)

        input_summary_text = package.get("input_summary_text")
        input_data_sample_html = package.get("input_data_sample_html")
        if input_summary_text or input_data_sample_html:
            input_html_parts = [f"<div><h3>Input Summary & Sample for '{html.escape(output_desc)}'</h3>"]
            if input_summary_text: input_html_parts.append(f"<p>{html.escape(input_summary_text)}</p>")
            if input_data_sample_html: input_html_parts.append(input_data_sample_html)
            input_html_parts.append("</div>")
            if len(input_html_parts) > 2:
                views.append(("html_rich", ''.join(input_html_parts), "Input Details"))

        step_logs = package.get("step_specific_logs")
        if step_logs and isinstance(step_logs, list):
            log_entries_html = "".join(
                f"<li><pre style='margin:0; padding:2px 0;'>{html.escape(log)}</pre></li>" for log in step_logs)
            log_html_full = f"<div><h3>Execution Logs for '{html.escape(output_desc)}'</h3><ul style='list-style-type:none; padding-left:5px; font-size:0.9em;'>{log_entries_html}</ul></div>"
            views.append(("html_rich", log_html_full, "Step Logs"))

        ## MODIFICATION: Add a view generator for the new period analysis package
        if output_type == "period_analysis_package":
            period_results = package.get("data", {})
            for period_name, period_data in period_results.items():
                # Add the plot for this period
                if period_data.get("figure_html"):
                    views.append(("plot_html", period_data["figure_html"], f"Figure: {period_name}"))
                # Add the data table for this period
                if period_data.get("table") is not None:
                    views.append(("table_df", period_data["table"], f"Data: {period_name}"))

        if output_type == "raw_df_full_summary":
            raw_df = package.get("raw_df_full")
            if raw_df is not None and isinstance(raw_df, pd.DataFrame):
                views.append(("table_df", raw_df, "Loaded Data Table"))
            data_html_preview = package.get("data_html")
            if data_html_preview:
                views.append(("html_rich", f"<div><h3>Data Preview</h3>{data_html_preview}</div>", "Data Preview"))

        elif output_type in ["plotly_html_summary", "plotly_graph_df", "plotly_graph", "plotly_graph_dict"]:
            # Static Image Page is now skipped for display as per user request
            # image_path = package.get("image_file_path")
            # if image_path and Path(image_path).exists():
            #     views.append(("image_file", image_path, "Figure (Static Image)"))

            plot_html_fragment = package.get("data_html_fragment")
            if plot_html_fragment:
                views.append(("plot_html", plot_html_fragment, "Figure (Interactive)"))

            raw_df_table = package.get("raw_df_table")
            if raw_df_table is not None and not raw_df_table.empty:
                views.append(("table_df", raw_df_table, "Associated Data Table"))

            data_summary_text = package.get("_data_summary_text")
            if data_summary_text:
                summary_html = f"<div><h3>Textual Summary of Figure</h3><p>{html.escape(data_summary_text)}</p></div>"
                views.append(("html_rich", summary_html, "Figure Summary"))

            if output_type == "plotly_graph_dict" and "keyword_clusters_data" in package:
                kw_clusters_data = package["keyword_clusters_data"]
                if isinstance(kw_clusters_data, dict):
                    summary_text = kw_clusters_data.get("summary_text", "No summary.")
                    cluster_details_json = json.dumps(kw_clusters_data.get("clusters", []), indent=2)
                    cluster_html = (
                        f"<div><h3>Keyword Cluster Analysis</h3><h4>Summary</h4><p>{html.escape(summary_text)}</p>"
                        f"<h4>Cluster Details (JSON)</h4><pre>{html.escape(cluster_details_json)}</pre></div>")
                    views.append(("html_rich", cluster_html, "Keyword Clusters"))

        elif output_type == "table_summary":
            raw_df_table = package.get("raw_df_table")
            if raw_df_table is not None and isinstance(raw_df_table, pd.DataFrame):
                views.append(("table_df", raw_df_table, "Data Table"))
            elif isinstance(output_data, str) and "<table" in output_data.lower():
                views.append(("html_rich", f"<div><h3>HTML Table</h3>{output_data}</div>", "HTML Table"))

        elif output_type == "html_section":
            ai_step_data_dict = output_data if isinstance(output_data, dict) else {}

            prompt_initial = ai_step_data_dict.get("initial_draft_prompt_sent", package.get("prompt_sent"))
            if prompt_initial and prompt_initial != "N/A for utility function":
                views.append(
                    ("html_rich", f"<div><h3>Initial Prompt</h3><pre>{html.escape(prompt_initial)}</pre></div>",
                     "AI Prompt (Initial)"))

            initial_draft_html = ai_step_data_dict.get("initial_draft_html_response")
            if initial_draft_html:
                views.append(
                    ("html_rich", f"<div><h3>Initial AI Draft</h3>{initial_draft_html}</div>", "AI Draft (Initial)"))

            r1_data_used = ai_step_data_dict.get("review_r1_data_used")
            if r1_data_used and isinstance(r1_data_used, dict):
                r1_html = f"<div><h3>Review Round 1 Data</h3><pre>{html.escape(json.dumps(r1_data_used, indent=2))}</pre></div>"
                views.append(("html_rich", r1_html, "AI Review R1 Data"))

            extracted_notes_pkg_key_candidate = f"{step_name}_ExtractedNotesPkg"  # Construct potential key
            # Check if self.completed_step_outputs is the correct place or if it's embedded in the current package
            extracted_notes_pkg = self.completed_step_outputs.get(extracted_notes_pkg_key_candidate,
                                                                  package.get("extracted_notes_package"))

            if extracted_notes_pkg and isinstance(extracted_notes_pkg, dict) and extracted_notes_pkg.get(
                    "type") == "coded_notes_list":
                notes_list = extracted_notes_pkg.get("data", [])
                notes_count = len(notes_list)
                notes_preview_html = "<ul>" + "".join(
                    f"<li><small><b>{html.escape(note.get('keyword_found', 'KW?'))}</b> (p.{html.escape(str(note.get('page_number', '?')))}): {html.escape(note.get('original_paragraph', '')[:80])}...</small></li>"
                    for note in notes_list[:5]) + "</ul>"
                notes_summary_html = f"<div><h3>Extracted Notes for Revision</h3><p>Number of notes: {notes_count}</p>{notes_preview_html if notes_list else ''}</div>"
                views.append(("html_rich", notes_summary_html, "AI Notes Context (R2)"))

            prompt_r2_text = ai_step_data_dict.get("prompt_sent_r2")
            if prompt_r2_text:
                views.append(
                    ("html_rich", f"<div><h3>Revision Prompt (R2)</h3><pre>{html.escape(prompt_r2_text)}</pre></div>",
                     "AI Prompt (Revision)"))

            response_html_content = ai_step_data_dict.get("response_html")
            if response_html_content:
                ai_output_title_suffix = "AI Generated Output"
                if initial_draft_html and response_html_content == initial_draft_html and not prompt_r2_text:
                    ai_output_title_suffix = "Initial AI Draft (View)"
                elif prompt_r2_text:
                    ai_output_title_suffix = "Final AI Section"
                elif package.get("description", "").startswith("AI Suggested"):
                    ai_output_title_suffix = "AI Suggestions"

                is_duplicate_view = any(
                    v[0] == "html_rich" and v[1] == response_html_content and v[2] == ai_output_title_suffix for v in
                    views)
                if not is_duplicate_view:
                    views.append(("html_rich",
                                  f"<div><h3>{html.escape(ai_output_title_suffix)}</h3>{response_html_content}</div>",
                                  ai_output_title_suffix))

            raw_ai_object = ai_step_data_dict.get("ai_response_raw_r2", ai_step_data_dict.get("ai_response_raw",
                                                                                              package.get(
                                                                                                  "ai_response_object")))
            if raw_ai_object:
                try:
                    raw_html = f"<div><h3>Raw AI Response Object</h3><pre>{html.escape(json.dumps(raw_ai_object, indent=2, default=str))}</pre></div>"
                except Exception:
                    raw_html = f"<div><h3>Raw AI Response Object</h3><pre>{html.escape(str(raw_ai_object))}</pre></div>"
                views.append(("html_rich", raw_html, "Raw AI Object"))

        elif output_type == "full_html_report":
            report_html_content = output_data if isinstance(output_data, str) else "Report content not available."
            views.append(
                ("html_rich", f"<div><h2>Full Report Preview</h2>{report_html_content}</div>", "Full Report Preview"))

        elif output_type in ["error", "interrupted", "no_output", "unknown_output",
                             "generic_data_dict"] or not output_type:
            data_str = ""
            try:
                data_str = json.dumps(output_data, indent=2, default=str) if output_data is not None else "N/A"
            except:
                data_str = str(output_data)

            status_html = (
                f"<div class='status-info-box' style='border:1px solid #ccc; padding:10px; margin-top:10px;'>"
                f"<h3>Status: {html.escape(output_desc)}</h3>"
                f"<p>Type: <code>{html.escape(str(output_type))}</code></p>"
                f"<p>Details:</p><pre>{html.escape(data_str)}</pre></div>")
            if output_type == "error": status_html = status_html.replace("status-info-box", "error-box").replace("#ccc",
                                                                                                                 "red")
            views.append(("html_rich", status_html, f"Status: {output_type.capitalize() if output_type else 'Info'}"))

        if not views and package:
            try:
                package_str = json.dumps(package, indent=2, default=lambda o: f"<non-serializable: {type(o).__name__}>")
            except Exception:
                package_str = str(package)[:2000] + "..."
            views.append(
                ("html_rich", f"<div><h3>Raw Step Output Package</h3><pre>{html.escape(package_str)}</pre></div>",
                 "Raw Package"))
        elif not views and not package:
            views.append(
                ("html_rich", f"<p>No data or package found for step {html.escape(step_name)}.</p>", "No Data"))

        if not views:
            views.append(("html_rich",
                          f"<p>No displayable output found for {html.escape(output_desc)} (Type: {html.escape(str(output_type))}).</p>",
                          "Empty Output"))

        return views
    def _find_step_list_item(self, step_tuple_to_find):
        target_phase, target_step_name = step_tuple_to_find
        for i in range(self.steps_progress_list.count()):
            item = self.steps_progress_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            if isinstance(item_data, dict) and \
                    item_data.get("phase") == target_phase and \
                    item_data.get("step") == target_step_name:
                return item
        return None

    def _on_step_list_item_clicked(self, item):
        item_data = item.data(Qt.ItemDataRole.UserRole)
        if not (isinstance(item_data, dict) and "phase" in item_data and "step" in item_data):
            return

        clicked_phase = item_data["phase"]
        clicked_step_name = item_data["step"]
        clicked_step_tuple = (clicked_phase, clicked_step_name)

        self.current_display_step_tuple = clicked_step_tuple
        self._handle_cumulative_log_message(f"Displaying output for: {clicked_step_name} ({clicked_phase})")

        # Case 1: Clicked step is the currently running step
        is_worker_active = self.worker_thread and self.worker_thread.isRunning()
        if clicked_step_tuple == (self._current_phase_running, self._current_step_running) and is_worker_active:
            self._handle_cumulative_log_message(
                f"Displaying live output for actively running step: {clicked_step_name}")
            if not self._current_step_live_buffer:
                # Initial state for a running step before any live data has arrived for it
                self.current_output_views = [("html_rich",
                                              f"<div><h3>Running: {html.escape(clicked_step_name)}</h3><p>Waiting for first live output...</p></div>",
                                              "Status - Running")]
            else:
                # Show the accumulated live buffer for the running step
                self.current_output_views = list(self._current_step_live_buffer)
            # Default to the latest live output page for the running step
            self.current_output_view_index = len(self.current_output_views) - 1 if self.current_output_views else 0

        # Case 2: Clicked step is a completed step with stored output
        elif clicked_step_tuple in self.completed_step_outputs:
            self._handle_cumulative_log_message(f"Displaying stored output for completed step: {clicked_step_name}")
            package = self.completed_step_outputs[clicked_step_tuple]
            self.current_output_views = self._generate_views_from_package(package, clicked_step_tuple)
            self.current_output_view_index = 0  # Default to the first page of the completed output

        # Case 3: Clicked step is pending, failed, or otherwise has no specific displayable content yet
        else:
            status = item_data.get("status", "unknown")
            status_message_detail = ""

            if status == "pending":
                if not (
                        self.worker_thread and self.worker_thread.isRunning() and self.analysis_worker and self.analysis_worker._is_running):  # Analysis not active
                    status_message_detail = "Analysis not started or has finished."
                else:  # Analysis active, step is pending
                    status_message_detail = "Waiting to be executed."
            elif status == "failed":
                status_message_detail = "Step failed. Check logs or previous step outputs for details."
            elif status == "interrupted":
                status_message_detail = "Step was interrupted."
            else:  # unknown, or other statuses
                status_message_detail = f"Output not yet available or step not run (Status: {html.escape(status.capitalize())})."

            self._handle_cumulative_log_message(f"Status for '{clicked_step_name}': {status_message_detail}")
            self.current_output_views = [("html_rich",
                                          f"<div><h3>{html.escape(clicked_step_name)}</h3><p>{html.escape(status_message_detail)}</p></div>",
                                          f"Status - {status.capitalize()}")]
            self.current_output_view_index = 0

        self._render_current_output_view()
        self._update_list_selection_from_display_tuple()  # Ensure the list highlights the clicked item
        self._update_navigation_button_states()  # Update main prev/next output buttons
        self._update_monitor_nav_buttons()  # Update monitor's own page prev/next buttons

    @pyqtSlot(str, str)
    def _handle_step_started(self, phase_name, step_name):
        self._current_phase_running = phase_name
        self._current_step_running = step_name
        self._current_step_live_buffer.clear()

        self.current_display_step_tuple = (phase_name, step_name)
        self.current_output_views.clear()  # no placeholder at all
        self.current_output_view_index = 0
        self._render_current_output_view()  # shows blank monitor until first log arrives

        # now record the “Starting …” message so it *is* the first live log line
        self._handle_cumulative_log_message(
            f"--- Starting: {step_name} ({phase_name}) ---"
        )
        self.current_output_view_index = 0
        self._render_current_output_view()



        self._update_navigation_button_states()

    @pyqtSlot(str)
    def _handle_cumulative_log_message(self, message):
        timestamp = QDateTime.currentDateTime().toString("hh:mm:ss")
        log_entry_text = f"[{timestamp}] {message}"
        self.log_monitor_browser.append(log_entry_text)

        if self._current_step_running and (self._current_phase_running,
                                           self._current_step_running) == self.current_display_step_tuple:
            log_view_content = f"<pre style='white-space:pre-wrap; word-wrap:break-word;'>{html.escape(log_entry_text)}</pre>"

            updated = False
            for i, (vtype, content, title) in enumerate(self._current_step_live_buffer):
                if vtype == "html_rich" and title == "Live Log":  # existing live-log view
                    self._current_step_live_buffer[i] = (
                        "html_rich",
                        content + log_view_content,  # extend html
                        "Live Log",
                    )
                    updated = True
                    break

            if not updated:  # none yet – replace placeholder or add new
                if (self._current_step_live_buffer and
                        "Waiting for live output..." in self._current_step_live_buffer[0][1]):
                    self._current_step_live_buffer[0] = ("html_rich", log_view_content, "Live Log")
                else:
                    self._current_step_live_buffer.append(("html_rich", log_view_content, "Live Log"))

    def _add_to_live_buffer_and_refresh(self, view_tuple: tuple):
        if not self._current_step_running: return

        if self._current_step_live_buffer and len(
                self._current_step_live_buffer) == 1 and "Waiting for live output..." in \
                self._current_step_live_buffer[0][1]:
            self._current_step_live_buffer[0] = view_tuple
        else:
            self._current_step_live_buffer.append(view_tuple)

        if (self._current_phase_running, self._current_step_running) == self.current_display_step_tuple:
            self.current_output_views = list(self._current_step_live_buffer)
            self.current_output_view_index = len(self.current_output_views) - 1
            self._render_current_output_view()

    @pyqtSlot(object, str)
    def _handle_live_table_data(self, dataframe, title_suffix):
        self._add_to_live_buffer_and_refresh(("table_df", dataframe.copy(), title_suffix))

    @pyqtSlot(object, str)
    def _handle_live_plotly_data(self, plotly_fig, title_suffix):
        html_plot_fragment = ""
        if plotly_fig:
            try:
                plotly_fig.update_layout(
                    template="plotly_dark" if self.current_theme_name == "dark" else "plotly_white")
                html_plot_fragment = plotly_fig.to_html(full_html=False, include_plotlyjs='cdn')
            except Exception as e:
                html_plot_fragment = f"<p>Error rendering live plot: {html.escape(str(e))}</p>"
        if html_plot_fragment:
            self._add_to_live_buffer_and_refresh(("plot_html", html_plot_fragment, title_suffix))

    @pyqtSlot(object, str)
    def _handle_live_html_data(self, data_object, title_suffix):
        html_content = ""
        if isinstance(data_object, dict):
            html_content = self._generate_rich_html_for_ai_object(data_object, title_suffix)
        elif isinstance(data_object, str):
            html_content = f"<div><h3>{html.escape(title_suffix)}</h3>{data_object}</div>"
        if html_content:
            self._add_to_live_buffer_and_refresh(("html_rich", html_content, title_suffix))

    @pyqtSlot(str, str)
    def _handle_live_image_data(self, image_path, title_suffix):
        if image_path and Path(image_path).exists():
            self._add_to_live_buffer_and_refresh(("image_file", image_path, title_suffix))

    @pyqtSlot(str, str, bool, object)
    def _handle_step_completed(self, phase_name, step_name, success, stored_output_package):
        icon = "✅" if success else "❌";
        color = QColor("green") if success else QColor("red")
        status_text = "completed"
        if stored_output_package and stored_output_package.get("type") == "error":
            status_text = "failed"
        elif stored_output_package and stored_output_package.get("type") == "interrupted":
            icon = "⏹️"; color = QColor("grey"); status_text = "interrupted"
        elif not success:
            status_text = "failed"

        self._handle_cumulative_log_message(f"--- {status_text.capitalize()}: {step_name} ({phase_name}) ---")
        current_step_tuple = (phase_name, step_name)

        if stored_output_package:
            self.completed_step_outputs[current_step_tuple] = stored_output_package
            if current_step_tuple not in self.executed_steps_chronological:
                self.executed_steps_chronological.append(current_step_tuple)

            is_worker_paused_for_manual = hasattr(self,
                                                  'analysis_worker') and self.analysis_worker and self.analysis_worker._is_paused and self.is_manual_mode

            if self.current_display_step_tuple == current_step_tuple and not is_worker_paused_for_manual:
                self.current_output_views = self._generate_views_from_package(stored_output_package, current_step_tuple)
                self.current_output_view_index = 0
                self._render_current_output_view()

        list_item = self._find_step_list_item(current_step_tuple)
        if list_item:
            list_item.setText(f"{icon} {step_name}")
            item_data = list_item.data(Qt.ItemDataRole.UserRole);
            item_data["status"] = status_text;
            list_item.setData(Qt.ItemDataRole.UserRole, item_data)
            list_item.setForeground(color)

        if (self._current_phase_running, self._current_step_running) == current_step_tuple:
            self._current_phase_running = None;
            self._current_step_running = None
            self._current_step_live_buffer.clear()

        self._update_navigation_button_states()

    def _render_current_output_view(self):
        if not self.current_output_views or not (0 <= self.current_output_view_index < len(self.current_output_views)):
            self.monitor_stack.setCurrentWidget(self.monitor_placeholder_widget)
            self.monitor_title_label.setText("Analysis Monitor")
            self._update_monitor_nav_buttons()
            return

        view_type, view_content, view_title_suffix = self.current_output_views[self.current_output_view_index]
        main_step_title = self.current_display_step_tuple[1] if self.current_display_step_tuple else "Output"
        if self.current_display_step_tuple and self.current_display_step_tuple in self.completed_step_outputs:
            main_step_title = self.completed_step_outputs[self.current_display_step_tuple].get("description",
                                                                                               main_step_title)

        self.monitor_title_label.setText(
            f"{main_step_title}: {view_title_suffix} ({self.current_output_view_index + 1}/{len(self.current_output_views)})")

        target_browser = self.html_content_browser
        current_widget_for_stack = self.html_content_browser

        if view_type == "plot_html":
            if WEBENGINE_AVAILABLE and isinstance(self.plotly_monitor_view,
                                                  QWebEngineView): target_browser = self.plotly_monitor_view
            html_to_set = view_content or "<p class='placeholder-text'>Plot data not available.</p>"
            if isinstance(target_browser, QWebEngineView):
                target_browser.setHtml(self._html_with_common_css(html_to_set), QUrl("about:blank"))
            else:
                target_browser.setHtml(self._html_with_common_css(html_to_set))
            current_widget_for_stack = target_browser
        elif view_type == "table_df":
            self._populate_table(view_content)
            current_widget_for_stack = self.table_monitor_widget
        elif view_type == "image_file":
            if WEBENGINE_AVAILABLE and isinstance(self.plotly_monitor_view,
                                                  QWebEngineView): target_browser = self.plotly_monitor_view
            image_path_str = view_content
            image_uri = QUrl.fromLocalFile(str(Path(image_path_str).resolve())).toString()
            image_html = f"<div style='padding:10px; text-align:center;'><img src='{html.escape(image_uri)}' alt='{html.escape(view_title_suffix)}' style='max-width:95%; max-height:85vh; object-fit:contain; border:1px solid #ccc;'/></div>"
            if isinstance(target_browser, QWebEngineView):
                target_browser.setHtml(self._html_with_common_css(image_html), QUrl("about:blank"))
            else:
                target_browser.setHtml(self._html_with_common_css(image_html))
            current_widget_for_stack = target_browser
        elif view_type == "html_rich":
            is_full_report_view = (self.current_display_step_tuple and self.current_display_step_tuple[
                1] == STEP_ASSEMBLE_REPORT_HTML and view_title_suffix == "Full Report Preview")
            if is_full_report_view and WEBENGINE_AVAILABLE and isinstance(self.plotly_monitor_view,
                                                                          QWebEngineView): target_browser = self.plotly_monitor_view

            html_to_set = view_content or "<p class='placeholder-text'>HTML content not available.</p>"
            if isinstance(target_browser, QWebEngineView):
                target_browser.setHtml(self._html_with_common_css(html_to_set), QUrl("about:blank"))
            else:
                target_browser.setHtml(self._html_with_common_css(html_to_set))
            current_widget_for_stack = target_browser
        else:
            self.monitor_placeholder_widget.setText(f"Unknown view type: {html.escape(view_type)}")
            current_widget_for_stack = self.monitor_placeholder_widget

        self.monitor_stack.setCurrentWidget(current_widget_for_stack)
        self._update_monitor_nav_buttons()

    def _on_previous_step_output_clicked(self):
        if not self.executed_steps_chronological: return
        current_idx = -1
        if self.current_display_step_tuple and self.current_display_step_tuple in self.executed_steps_chronological:
            try:
                current_idx = self.executed_steps_chronological.index(self.current_display_step_tuple)
            except ValueError:
                pass

        if current_idx > 0:
            prev_step_tuple = self.executed_steps_chronological[current_idx - 1]
            item_to_click = self._find_step_list_item(prev_step_tuple)
            if item_to_click: self._on_step_list_item_clicked(item_to_click)
        elif self.executed_steps_chronological:
            first_step_tuple = self.executed_steps_chronological[0]
            item_to_click = self._find_step_list_item(first_step_tuple)
            if item_to_click: self._on_step_list_item_clicked(item_to_click)

    def _on_next_step_output_clicked(self):
        is_running_live = self.worker_thread and self.worker_thread.isRunning()
        is_paused_by_worker = hasattr(self,
                                      'analysis_worker') and self.analysis_worker and self.analysis_worker._is_paused

        if self.is_manual_mode and is_running_live and is_paused_by_worker:
            self._handle_cumulative_log_message("Proceeding to execute next manual step...")
            if self.analysis_worker: self.analysis_worker.resume_manual_step()
        elif self.executed_steps_chronological:
            current_idx = -1
            if self.current_display_step_tuple and self.current_display_step_tuple in self.executed_steps_chronological:
                try:
                    current_idx = self.executed_steps_chronological.index(self.current_display_step_tuple)
                except ValueError:
                    pass

            if current_idx < len(self.executed_steps_chronological) - 1:
                next_step_tuple = self.executed_steps_chronological[current_idx + 1]
                item_to_click = self._find_step_list_item(next_step_tuple)
                if item_to_click: self._on_step_list_item_clicked(item_to_click)
            elif self.executed_steps_chronological:
                last_step_tuple = self.executed_steps_chronological[-1]
                item_to_click = self._find_step_list_item(last_step_tuple)
                if item_to_click: self._on_step_list_item_clicked(item_to_click)
        self._update_navigation_button_states()

    def _populate_ai_provider_combo(self):
        self.ai_provider_combo.clear()
        self.ai_provider_combo.addItems(AI_MODELS_CONFIG.keys())
        if self.current_ai_provider_key in AI_MODELS_CONFIG:
            self.ai_provider_combo.setCurrentText(self.current_ai_provider_key)
        elif self.ai_provider_combo.count() > 0:
            self.ai_provider_combo.setCurrentIndex(
                0); self.current_ai_provider_key = self.ai_provider_combo.currentText()
        self._update_ai_model_combo()

    def _update_ai_model_combo(self):
        provider_display_name = self.ai_provider_combo.currentText();
        self.ai_model_combo.clear()
        if provider_display_name and provider_display_name in AI_MODELS_CONFIG:
            models = AI_MODELS_CONFIG[provider_display_name];
            model_display_names = list(models.keys());
            self.ai_model_combo.addItems(model_display_names)
            found_current_model = False
            for display_name, api_name in models.items():
                if api_name == self.current_ai_model_api_name: self.ai_model_combo.setCurrentText(
                    display_name); found_current_model = True; break
            if not found_current_model and model_display_names: self.ai_model_combo.setCurrentIndex(0)
        self._on_ai_model_changed()

    def _on_ai_provider_changed(self):
        self.current_ai_provider_key = self.ai_provider_combo.currentText(); self._update_ai_model_combo()

    def _on_ai_model_changed(self):
        provider_display_name = self.ai_provider_combo.currentText();
        model_display_name = self.ai_model_combo.currentText()
        if provider_display_name in AI_MODELS_CONFIG and model_display_name in AI_MODELS_CONFIG[provider_display_name]:
            self.current_ai_provider_key = provider_display_name;
            self.current_ai_model_api_name = AI_MODELS_CONFIG[provider_display_name][model_display_name]
            logging.debug(
                f"AI Config Updated: Provider='{self.current_ai_provider_key}', Model API='{self.current_ai_model_api_name}'")
        elif AI_MODELS_CONFIG and self.ai_provider_combo.count() > 0 and self.ai_model_combo.count() > 0:
            self.ai_provider_combo.setCurrentIndex(0);
            self.current_ai_provider_key = self.ai_provider_combo.currentText()
            models_for_first_provider = AI_MODELS_CONFIG[self.current_ai_provider_key]
            self.ai_model_combo.setCurrentIndex(0);
            first_model_display = self.ai_model_combo.currentText()
            self.current_ai_model_api_name = models_for_first_provider[first_model_display]
            logging.warning(
                f"AI Model selection issue, reset to default: {self.current_ai_provider_key} - {self.current_ai_model_api_name}")
        else:

            self._harden_ai_model_selection()
            return

    def _initiate_data_load(self, source_type):
        self.current_data_source_type = source_type;
        self.current_data_file_path = None;
        self.current_zotero_collection_name = None
        self.completed_step_outputs.clear();
        self.executed_steps_chronological.clear();
        self.current_display_step_tuple = None
        if source_type == "file":
            file_path, _ = QFileDialog.getOpenFileName(self, "Open Bibliographic Data File", "",
                                                       "Data Files (*.csv *.xlsx *.xls);;All Files (*)")
            if not file_path: self.data_status_label.setText(
                "File loading cancelled."); self.run_analysis_btn.setEnabled(False); return
            self.current_data_file_path = file_path;
            self.data_status_label.setText(f"File: {Path(file_path).name}. Ready.")
        elif source_type == "zotero":
            self.current_zotero_collection_name = self.zotero_collection_input.text().strip()
            coll_name_disp = f"'{self.current_zotero_collection_name}'" if self.current_zotero_collection_name else "All Library Items"
            self.data_status_label.setText(f"Zotero: {coll_name_disp}. Ready.")
        self.run_analysis_btn.setEnabled(True);
        self._clear_monitor_and_steps_list();
        self._populate_steps_list_statically()

    def _on_analysis_type_changed(self, analysis_type):
        self._populate_steps_list_statically(); self.completed_step_outputs.clear(); self.executed_steps_chronological.clear(); self.current_display_step_tuple = None; self._update_navigation_button_states()

    def _populate_steps_list_statically(self):
        self.steps_progress_list.clear();
        analysis_type = self.analysis_type_combo.currentText()
        try:
            if 'BibliographicAnalysisWorker' in globals():
                dummy_worker = BibliographicAnalysisWorker("", "", "", "", "", False, "", "", "")
                ## MODIFICATION: Get step lists from all available configs in the dummy worker
                full_review_research_steps = [s[0] for s in dummy_worker.full_review_steps_config.get(PHASE_RESEARCH, [])]
                full_review_writing_steps = [s[0] for s in dummy_worker.full_review_steps_config.get(PHASE_WRITING, [])]
                chrono_analysis_steps = [s[0] for s in dummy_worker.chronological_review_steps_config.get(PHASE_CHRONO_ANALYSIS, [])]
                chrono_writing_steps = [s[0] for s in dummy_worker.chronological_review_steps_config.get(PHASE_CHRONO_WRITING, [])]
                chrono_synthesis_steps = [s[0] for s in dummy_worker.chronological_review_steps_config.get(PHASE_CHRONO_SYNTHESIS, [])]
                del dummy_worker
            else:
                raise NameError("BibliographicAnalysisWorker not defined globally for step list population.")
        except Exception as e_dummy:
            # Fallback remains the same, as it's just for emergencies.
            logging.error(
                f"Error accessing worker config for step list: {e_dummy}. Using hardcoded fallback for steps.")
            # ... (fallback step lists are unchanged)
            return

        ## MODIFICATION: Select the correct step list based on the chosen analysis type
        if analysis_type == ANALYSIS_TYPE_FULL_REVIEW:
            static_display_config = {
                PHASE_RESEARCH: full_review_research_steps,
                PHASE_WRITING: full_review_writing_steps
            }
        elif analysis_type == ANALYSIS_TYPE_CHRONOLOGICAL_REVIEW:
            static_display_config = {
                PHASE_CHRONO_ANALYSIS: chrono_analysis_steps,
                PHASE_CHRONO_WRITING: chrono_writing_steps,
                PHASE_CHRONO_SYNTHESIS: chrono_synthesis_steps
            }
        else: # Single analysis types
            self.steps_progress_list.setVisible(True);
            item = QListWidgetItem(f"⏳ {analysis_type}")
            single_analysis_phase = PHASE_RESEARCH
            item.setData(Qt.ItemDataRole.UserRole,
                         {"phase": single_analysis_phase, "step": analysis_type, "status": "pending"});
            self.steps_progress_list.addItem(item)
            return

        self.steps_progress_list.setVisible(True)
        for phase_name, step_names_in_phase in static_display_config.items():
            phase_item = QListWidgetItem(f"--- {phase_name} ---");
            phase_item.setFlags(phase_item.flags() & ~Qt.ItemFlag.ItemIsSelectable);
            font = phase_item.font();
            font.setBold(True);
            phase_item.setFont(font);
            self.steps_progress_list.addItem(phase_item)
            for step_name in step_names_in_phase:
                item = QListWidgetItem(f"⏳ {step_name}");
                item.setData(Qt.ItemDataRole.UserRole,
                             {"phase": phase_name, "step": step_name, "status": "pending"});
                self.steps_progress_list.addItem(item)


    def _on_mode_changed(self, checked):
        self.is_manual_mode = self.manual_mode_radio.isChecked()
        if self.analysis_worker:
            try:
                self.analysis_worker.is_manual_mode = self.is_manual_mode
            except RuntimeError:
                logging.warning("Set mode on deleted worker.")
        self._update_navigation_button_states()

    def _on_pause_resume_clicked(self):
        if not self.analysis_worker or not self.worker_thread: self.pause_resume_btn.setChecked(
            False); self._update_navigation_button_states(); return
        try:
            if not self.worker_thread.isRunning(): self.pause_resume_btn.setChecked(
                False); self._update_navigation_button_states(); return
        except RuntimeError:
            self.pause_resume_btn.setChecked(False); self._update_navigation_button_states(); return
        is_now_paused = self.analysis_worker.toggle_pause_resume()
        self.pause_resume_btn.setChecked(is_now_paused)
        self._update_navigation_button_states()

    def _on_stop_analysis_clicked(self):
        if self.analysis_worker:
            try:
                self._handle_cumulative_log_message("Stop request sent."); self.analysis_worker.request_stop()
            except RuntimeError:
                logging.warning("Attempted to stop a deleted worker.")
        if self.worker_thread:
            try:
                self.worker_thread.quit(); self.worker_thread.wait(3000)
            except RuntimeError:
                logging.warning("Worker thread already gone or error quitting.")
        self.analysis_worker = None;
        self.worker_thread = None
        self._current_phase_running = None;
        self._current_step_running = None;
        self._current_step_live_buffer.clear()
        self._update_analysis_controls_state(is_running=False)

    @pyqtSlot()
    def _clear_worker_references(self):
        logging.debug("Worker thread finished signal received. Clearing references if any remain.")
        self.analysis_worker = None;
        self.worker_thread = None
        self._current_phase_running = None;
        self._current_step_running = None;
        self._current_step_live_buffer.clear()
        self._update_analysis_controls_state(is_running=False)

    @pyqtSlot(bool, str)
    @pyqtSlot(bool, str)
    def _handle_analysis_finished(self, success, message):
        self._handle_cumulative_log_message(f"===== ANALYSIS FINISHED: {message} (Success: {success}) =====")
        self.monitor_title_label.setText(f"Analysis Finished: {message}")

        # --- stop/clear the worker thread (unchanged behavior) ---
        if self.worker_thread is not None:
            try:
                if self.worker_thread.isRunning():
                    self.worker_thread.quit()
            except RuntimeError:
                self.worker_thread = None
                self._clear_worker_references()
        else:
            self._clear_worker_references()

        # --- Batch follow-up: if toggle is ON and we just did store-only, run the batch then re-run in READ mode ---
        try:
            import os, re
            from pathlib import Path

            # Toggle must exist & be ON
            if not (hasattr(self, "batch_mode_checkbox") and self.batch_mode_checkbox.isChecked()):
                return

            # If we're already in READ mode, don't loop
            if os.getenv("BATCH_MODE_READ", "0") == "1":
                return

            # Determine topic (used as function/collection names)
            topic = getattr(self, "current_research_topic", None) or "Default_Research_Topic"

            # Sanitize for filesystem
            def _sanitize(s: str) -> str:
                return re.sub(r"[^A-Za-z0-9_.-]+", "_", (s or "task")).strip("_")

            eff_function = _sanitize(topic)
            eff_collection = _sanitize(topic)

            # Process the batch for this function/collection
            self._handle_cumulative_log_message(
                f"Batch mode: submitting/polling batch for function='{eff_function}', collection='{eff_collection}'…"
            )

            ok = _process_batch_for(function=eff_function, collection_name=eff_collection)
            if not ok:
                self._handle_cumulative_log_message(
                    "Batch processing did not complete successfully; skipping READ re-run.")
                return

            # Flip env so downstream call_models() runs in READ mode
            os.environ["BATCH_FUNCTION"] = topic
            os.environ["BATCH_COLLECTION"] = topic
            os.environ["BATCH_MODE_STORE_ONLY"] = "0"
            os.environ["BATCH_MODE_READ"] = "1"

            self._handle_cumulative_log_message("Batch output ready. Re-running in READ mode to ingest results…")
            QTimer.singleShot(0, self._start_analysis_process)

        except Exception as e:
            # Never crash the UI on batch follow-up
            try:
                self._handle_cumulative_log_message(f"Batch follow-up skipped due to error: {e}")
            except Exception:
                pass

    @pyqtSlot(str, str)
    def _handle_worker_paused_manual(self, phase_name, step_name):
        self._handle_cumulative_log_message(f"--- PAUSED (Manual Mode) after: {step_name} ({phase_name}) ---")
        self._handle_cumulative_log_message("Click 'Execute Next' to proceed, or navigate outputs.")
        self.monitor_title_label.setText(f"Paused after: {step_name} (Manual Mode)")
        self.pause_resume_btn.setChecked(True)
        self._update_navigation_button_states()

    def _update_list_selection_from_display_tuple(self):
        if not self.current_display_step_tuple: self.steps_progress_list.clearSelection(); return
        target_phase, target_step = self.current_display_step_tuple
        list_item = self._find_step_list_item((target_phase, target_step))
        if list_item:
            self.steps_progress_list.setCurrentItem(list_item)
        else:
            self.steps_progress_list.clearSelection()

    def _update_navigation_button_states(self):
        is_actively_processing = False
        is_worker_paused_by_toggle = False
        is_worker_paused_for_manual_step = False

        if self.worker_thread and self.worker_thread.isRunning() and hasattr(self,
                                                                             'analysis_worker') and self.analysis_worker:
            try:
                is_actively_processing = self.analysis_worker._is_running
                is_worker_paused_by_toggle = self.analysis_worker._is_paused
                if self.is_manual_mode and self.analysis_worker._is_paused:  # Worker is paused specifically for a manual step
                    is_worker_paused_for_manual_step = True
                # If not in manual mode, any pause is a user toggle pause
                if not self.is_manual_mode and self.analysis_worker._is_paused:
                    is_worker_paused_by_toggle = True  # Explicitly set for clarity
                else:  # If in manual mode, but not paused for a manual step, it's not a toggle pause
                    if self.is_manual_mode and not is_worker_paused_for_manual_step:
                        is_worker_paused_by_toggle = False


            except RuntimeError:  # Worker might be gone
                is_actively_processing = False
                is_worker_paused_by_toggle = False
                is_worker_paused_for_manual_step = False

        # Pause/Resume button logic
        self.pause_resume_btn.setEnabled(is_actively_processing and not is_worker_paused_for_manual_step)
        self.pause_resume_btn.setChecked(is_worker_paused_by_toggle and not is_worker_paused_for_manual_step)
        if is_worker_paused_by_toggle and not is_worker_paused_for_manual_step:
            self.pause_resume_btn.setText("Resume")
            self.pause_resume_btn.setIcon(QIcon.fromTheme("media-playback-start"))
        else:
            self.pause_resume_btn.setText("Pause")
            self.pause_resume_btn.setIcon(QIcon.fromTheme("media-playback-pause"))

        # Previous Output Button
        can_go_prev_output = False
        if self.executed_steps_chronological and self.current_display_step_tuple:
            try:
                current_idx = self.executed_steps_chronological.index(self.current_display_step_tuple)
                can_go_prev_output = current_idx > 0
            except ValueError:  # current_display_step_tuple might not be in executed_steps (e.g. live view)
                can_go_prev_output = bool(self.executed_steps_chronological)
                # Allow navigating stored outputs freely.
        self.prev_step_btn.setEnabled(can_go_prev_output)

        # Next Output / Execute Next Button
        can_execute_next_manual = self.is_manual_mode and is_worker_paused_for_manual_step
        can_view_next_completed_output = False
        if self.executed_steps_chronological and self.current_display_step_tuple:
            try:
                current_idx = self.executed_steps_chronological.index(self.current_display_step_tuple)
                can_view_next_completed_output = current_idx < len(self.executed_steps_chronological) - 1
            except ValueError:
                # If current_display_step_tuple is not in executed_steps (e.g., it's a live view of a running step,
                # or a pending step), then "Next Output" should be enabled if there are *any* executed steps
                # that could potentially be "next" relative to the start or a known position.
                # A simpler approach: enable if there are more completed steps to see.
                if self.executed_steps_chronological:
                    # If current_display_step_tuple is not a completed one, effectively treat current_idx as -1
                    # so that it can go to the first completed one if it exists.
                    # Or, if current_display_step_tuple IS a completed one, this works as before.
                    # This logic is complex; the primary navigation is the list.
                    # For now, stick to enabling if there's a defined "next" completed output.
                    pass  # Keep original logic for can_view_next_completed_output

        if can_execute_next_manual:
            self.next_step_btn.setEnabled(True)
            self.next_step_btn.setText("Execute Next")
            self.next_step_btn.setIcon(QIcon.fromTheme("media-skip-forward"))  # Or a more fitting "execute" icon
        else:
            # Allow navigating stored outputs freely.
            self.next_step_btn.setEnabled(can_view_next_completed_output)
            self.next_step_btn.setText("Next Output")
            self.next_step_btn.setIcon(QIcon.fromTheme("go-next"))

        self.stop_analysis_btn.setEnabled(is_actively_processing)

    def _populate_table(self, df):
        self.table_monitor_widget.clearContents();
        self.table_monitor_widget.setRowCount(0);
        self.table_monitor_widget.setColumnCount(0)
        if df is None or df.empty: self._handle_cumulative_log_message("No data for table."); return
        if len(df) > 1000: logging.warning(f"Populating table with {len(df)} rows.");
        self.table_monitor_widget.setRowCount(len(df));
        self.table_monitor_widget.setColumnCount(len(df.columns));
        self.table_monitor_widget.setHorizontalHeaderLabels(df.columns.astype(str))
        for i, row_tuple in enumerate(df.itertuples(index=False)):
            for j, value in enumerate(row_tuple): self.table_monitor_widget.setItem(i, j, QTableWidgetItem(
                str(value) if pd.notna(value) else ""))
        self.table_monitor_widget.resizeColumnsToContents()
        header = self.table_monitor_widget.horizontalHeader()
        if df.shape[1] > 4:
            for j_col in range(df.shape[1] - 1): header.setSectionResizeMode(j_col,
                                                                             QHeaderView.ResizeMode.ResizeToContents)
            header.setSectionResizeMode(df.shape[1] - 1, QHeaderView.ResizeMode.Stretch)
        else:
            header.setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

    def _clear_monitor_and_steps_list(self):
        self.log_monitor_browser.clear()
        if WEBENGINE_AVAILABLE and isinstance(self.plotly_monitor_view, QWebEngineView):
            self.plotly_monitor_view.setHtml(self._html_with_common_css("<p>Monitor cleared.</p>"))
        elif isinstance(self.plotly_monitor_view, QTextBrowser):
            self.plotly_monitor_view.clear()
        self.table_monitor_widget.clearContents();
        self.table_monitor_widget.setRowCount(0)
        self.html_content_browser.setHtml(self._html_with_common_css("<p>Monitor cleared.</p>"))
        self.monitor_stack.setCurrentWidget(self.monitor_placeholder_widget)
        self.monitor_title_label.setText("Analysis Monitor")
        self.steps_progress_list.clear()
        self.current_output_views.clear()
        self.current_output_view_index = 0
        self._update_monitor_nav_buttons()

    def _update_analysis_controls_state(self, is_running=False):
        self.load_from_zotero_btn.setEnabled(not is_running);
        self.zotero_collection_input.setEnabled(not is_running);
        self.load_file_btn.setEnabled(not is_running)
        can_initiate_run = hasattr(self, 'current_data_source_type') and self.current_data_source_type is not None
        self.run_analysis_btn.setEnabled(can_initiate_run and not is_running)
        self.analysis_type_combo.setEnabled(not is_running);
        self.manual_mode_radio.setEnabled(not is_running);
        self.auto_mode_radio.setEnabled(not is_running)
        self.ai_provider_combo.setEnabled(not is_running);
        self.ai_model_combo.setEnabled(not is_running)
        ## MODIFICATION: Check for either report step name when enabling export buttons
        full_report_step_tuple = (PHASE_WRITING, STEP_ASSEMBLE_REPORT_HTML)
        chrono_report_step_tuple = (PHASE_CHRONO_SYNTHESIS, STEP_ASSEMBLE_REPORT_HTML)

        report_generated_successfully = (
            (full_report_step_tuple in self.completed_step_outputs and
             self.completed_step_outputs[full_report_step_tuple].get("type") == "full_html_report")
            or
            (chrono_report_step_tuple in self.completed_step_outputs and
             self.completed_step_outputs[chrono_report_step_tuple].get("type") == "full_html_report")
        )

        self.export_html_btn.setEnabled(report_generated_successfully and not is_running)
        self.export_pdf_btn.setEnabled(report_generated_successfully and not is_running)
        self.export_docx_btn.setEnabled(report_generated_successfully and not is_running)
        self._update_navigation_button_states()

    def closeEvent(self, event):
        thread_is_running = False
        if self.worker_thread is not None:
            try:
                thread_is_running = self.worker_thread.isRunning()
            except RuntimeError:
                thread_is_running = False; self.worker_thread = None
        if thread_is_running:
            reply = QMessageBox.question(self, "Analysis in Progress",
                                         "An analysis is currently running. Are you sure you want to close?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                         QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                if self.analysis_worker:
                    try:
                        self.analysis_worker.request_stop()
                    except RuntimeError:
                        logging.warning("Stop deleted worker during close.")
                if self.worker_thread:
                    try:
                        self.worker_thread.quit(); self.worker_thread.wait(1000)
                    except RuntimeError:
                        pass
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()

    def apply_theme(self, theme_name="dark"):
        self.current_theme_name = theme_name;
        is_dark = theme_name == "dark"
        # Using CSS variables defined in COMMON_MONITOR_CSS for consistency
        self.setStyleSheet(f"""
              QWidget#bibliographicAnalysisWidget, QFrame#analysisControlsPanel, QFrame#analysisMonitorPanel {{ 
                  background-color: var(--bg-{theme_name}); 
                  color: var(--fg-{theme_name});
              }}
              QGroupBox {{ 
                  border: 1px solid var(--border-color-{theme_name}); 
                  border-radius: 5px; margin-top: 6px; padding: 8px; 
                  background-color: {"#2b2b2b" if is_dark else "#f0f4f8"}; /* More specific than general var */
              }}
              QGroupBox::title {{ 
                  subcontrol-origin: margin; subcontrol-position: top left; padding: 0 3px; 
                  font-weight: bold; color: var(--fg-{theme_name});
              }}
              QLabel, QRadioButton {{ color: var(--fg-{theme_name}); }} 
              QLabel#dataStatusLabel {{ color: {"#A0A0A0" if is_dark else "#707070"}; font-style: italic; }}
              QLineEdit, QComboBox, QListWidget, QTextBrowser#logMonitorBrowser {{ 
                  background-color: var(--pre-bg-{theme_name}); /* Using pre-bg for inputs */
                  color: var(--fg-{theme_name}); 
                  border: 1px solid var(--border-color-{theme_name}); 
                  border-radius: 3px; padding: 4px; 
              }}
              QTableWidget {{ 
                  background-color: var(--bg-{theme_name}); 
                  color: var(--fg-{theme_name}); 
                  border: 1px solid var(--border-color-{theme_name}); 
                  gridline-color: var(--border-color-{theme_name}); 
              }}
              QTableWidget QHeaderView::section {{ 
                  background-color: var(--header-bg-{theme_name}); 
                  color: var(--fg-{theme_name});
                  padding: 6px; 
                  border: 1px solid var(--border-color-{theme_name});
                  font-weight: bold;
              }}
               QTableWidget::item {{ padding: 5px; }}
               QTableWidget::item:alternate {{ background-color: {"#2d333b" if is_dark else "#f6f8fa"}; }} /* Alternate row color */
               QTableWidget::item:selected {{ background-color: {"#0078D7" if not is_dark else "#005a9e"}; color: white; }}


              QPushButton {{ 
                  background-color: {"#505070" if is_dark else "#D0D0E0"}; 
                  color: var(--fg-{theme_name}); 
                  border: 1px solid var(--border-color-{theme_name}); 
                  padding: 6px 10px; border-radius: 3px; 
              }}
              QPushButton:hover {{ background-color: {"#606080" if is_dark else "#C0C0D0"}; }} 
              QPushButton:disabled {{ background-color: {"#404040" if is_dark else "#E0E0E0"}; color: grey; }}
              QPushButton#runAnalysisBtn {{ font-weight: bold; background-color: {"#4CAF50" if is_dark else "#5CB85C"}; color: white; }} 
              QPushButton#runAnalysisBtn:hover {{ background-color: {"#45a049" if is_dark else "#4CAE4C"}; }}
              QListWidget::item:selected {{ background-color: {"#0078D7" if not is_dark else "#005a9e"}; color: white; }}
          """)

        input_bg_qcolor = QColor("#252525") if is_dark else QColor("white")  # For palette-based widgets
        text_c_qcolor = QColor("#E0E0E0") if is_dark else QColor("#1E1E1E")

        if WEBENGINE_AVAILABLE and isinstance(self.plotly_monitor_view,
                                              QWebEngineView) and self.plotly_monitor_view.page():
            self.plotly_monitor_view.page().setBackgroundColor(input_bg_qcolor)
        elif isinstance(self.plotly_monitor_view, QTextBrowser):
            p = self.plotly_monitor_view.palette();
            p.setColor(QPalette.ColorRole.Base, input_bg_qcolor);
            p.setColor(QPalette.ColorRole.Text, text_c_qcolor);
            self.plotly_monitor_view.setPalette(p)

        p_html = self.html_content_browser.palette();
        p_html.setColor(QPalette.ColorRole.Base, input_bg_qcolor);
        p_html.setColor(QPalette.ColorRole.Text, text_c_qcolor);
        self.html_content_browser.setPalette(p_html)

        # Explicitly set log_monitor_browser palette if not covered by general QTextBrowser style
        p_log = self.log_monitor_browser.palette();
        p_log.setColor(QPalette.ColorRole.Base, input_bg_qcolor);
        p_log.setColor(QPalette.ColorRole.Text, text_c_qcolor);
        self.log_monitor_browser.setPalette(p_log)

        if self.current_output_views: self._render_current_output_view()  # Re-render with new theme
        self.update()
    def _get_full_report_html_content(self):
        ## MODIFICATION: Check both possible locations for a completed report
        full_report_step_tuple = (PHASE_WRITING, STEP_ASSEMBLE_REPORT_HTML)
        chrono_report_step_tuple = (PHASE_CHRONO_SYNTHESIS, STEP_ASSEMBLE_REPORT_HTML)
        report_package = None

        if full_report_step_tuple in self.completed_step_outputs:
            report_package = self.completed_step_outputs[full_report_step_tuple]
        elif chrono_report_step_tuple in self.completed_step_outputs:
            report_package = self.completed_step_outputs[chrono_report_step_tuple]

        if report_package and report_package.get("type") == "full_html_report":
            return report_package.get("data")

        self._handle_cumulative_log_message("Full HTML report not yet generated or available for export.")
        QMessageBox.warning(self, "Report Not Ready", "The full HTML report has not been generated yet.")
        return None

    @pyqtSlot()
    def _export_report_as_html(self):
        """
        Save the assembled report as a standalone UTF-8 HTML file.
        """
        html_content = self._get_full_report_html_content()
        if not html_content:
            return
        default_filename = f"Bibliometric_Report_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}.html"
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save HTML Report", default_filename, "HTML Files (*.html *.htm)"
        )
        if not filePath:
            self._handle_cumulative_log_message("Export HTML cancelled by user.")
            return
        try:
            with open(filePath, "w", encoding="utf-8") as f:
                f.write(html_content)
            self._handle_cumulative_log_message(f"Report successfully saved as HTML: {filePath}")
            QMessageBox.information(self, "Export Successful", f"Report saved to:\n{filePath}")
        except Exception as e:
            self._handle_cumulative_log_message(f"Error saving HTML report: {e}")
            QMessageBox.critical(self, "Export Error", f"Could not save HTML report:\n{e}")

    @pyqtSlot()
    def _export_report_as_pdf(self):
        """
        Export the assembled report to PDF.
        Path A: PyQtWebEngine's printToPdf (preserves Plotly via HTML rendering).
        Path B: Fallback to Qt printer if WebEngine is unavailable.
        """
        html_content = self._get_full_report_html_content()
        if not html_content:
            return

        # Primary path: WebEngine rendering -> printToPdf
        if WEBENGINE_AVAILABLE:
            try:
                default_filename = f"Bibliometric_Report_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}.pdf"
                filePath, _ = QFileDialog.getSaveFileName(self, "Save PDF Report", default_filename,
                                                          "PDF Files (*.pdf)")
                if not filePath:
                    self._handle_cumulative_log_message("Export PDF cancelled by user.")
                    return

                # Fresh headless view for printing to avoid touching the main monitor view
                view = QWebEngineView()

                def _on_loaded(ok: bool):
                    if not ok:
                        QMessageBox.critical(self, "PDF Export Error", "Failed to load HTML content for PDF export.")
                        view.deleteLater()
                        return
                    view.page().printToPdf(filePath)
                    self._handle_cumulative_log_message(f"PDF export initiated to: {filePath}.")
                    QMessageBox.information(self, "PDF Export", f"PDF saved to:\n{filePath}")
                    view.deleteLater()

                view.loadFinished.connect(_on_loaded)
                view.setHtml(html_content, QUrl("about:blank"))
                return
            except Exception as e:
                self._handle_cumulative_log_message(f"PDF export via WebEngine failed: {e}")

        # Fallback: Qt printer route (no WebEngine)
        try:
            from PyQt6.QtPrintSupport import QPrinter
            from PyQt6.QtGui import QTextDocument
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            default_filename = f"Bibliometric_Report_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}.pdf"
            filePath, _ = QFileDialog.getSaveFileName(self, "Save PDF Report", default_filename, "PDF Files (*.pdf)")
            if not filePath:
                self._handle_cumulative_log_message("Export PDF cancelled by user.")
                return
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(filePath)
            doc = QTextDocument()
            doc.setHtml(html_content)
            doc.print(printer)
            self._handle_cumulative_log_message(f"Export PDF (Qt printer): saved to {filePath}")
            QMessageBox.information(self, "Export Successful", f"Report saved to:\n{filePath}")
        except Exception as e:
            self._handle_cumulative_log_message(f"Qt printer PDF export failed: {e}")
            QMessageBox.critical(self, "Export Error", f"Could not save PDF report:\n{e}")

    @pyqtSlot()
    def _export_report_as_docx(self):
        """
        Export the assembled report to DOCX via Pandoc (preferred).
        Falls back to a minimal python-docx conversion if Pandoc is unavailable.
        """
        html_content = self._get_full_report_html_content()
        if not html_content:
            return

        default_filename = f"Bibliometric_Report_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}.docx"
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save DOCX Report", default_filename, "Word Documents (*.docx)"
        )
        if not filePath:
            self._handle_cumulative_log_message("Export DOCX cancelled by user.")
            return

        # Attempt 1: Pandoc (best fidelity)
        try:
            subprocess.run(['pandoc', '--version'], check=True, capture_output=True)
            temp_html = None
            try:
                tmp = QTemporaryFile(str(Path(QDir.tempPath()) / "temp_report_XXXXXX.html"))
                tmp.setAutoRemove(False)
                if not tmp.open():
                    raise IOError("Could not open temporary file for DOCX conversion.")
                temp_html = tmp.fileName()
                with open(temp_html, "w", encoding="utf-8") as f:
                    f.write(html_content)
                tmp.close()

                result = subprocess.run(
                    ['pandoc', temp_html, '-s', '--wrap=auto', '--embed-resources', '--standalone', '-o',
                     str(filePath)],
                    capture_output=True, text=True, check=False
                )
                if result.returncode == 0:
                    self._handle_cumulative_log_message(f"Report saved as DOCX: {filePath}")
                    QMessageBox.information(self, "Export Successful", f"Report saved to:\n{filePath}")
                    return
                else:
                    self._handle_cumulative_log_message(f"Pandoc error (Code {result.returncode}): {result.stderr}")
                    QMessageBox.critical(self, "DOCX Export Error",
                                         f"Pandoc conversion failed:\n{result.stderr[:600]}.")
            finally:
                if temp_html and Path(temp_html).exists():
                    try:
                        os.remove(temp_html)
                    except Exception as e_del:
                        logging.warning(f"Could not delete temp DOCX HTML file {temp_html}: {e_del}")
        except Exception as e:
            self._handle_cumulative_log_message(f"Pandoc not available or failed pre-check: {e}")

        # Attempt 2: Minimal python-docx conversion
        try:
            from bs4 import BeautifulSoup
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            soup = BeautifulSoup(html_content, "html.parser")
            doc = Document()

            title = soup.find("h1")
            if title and title.get_text(strip=True):
                p = doc.add_paragraph(title.get_text(strip=True))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(18)

            for tag in soup.find_all(["h2", "h3", "h4", "p", "li"]):
                text = tag.get_text(" ", strip=True)
                if not text:
                    continue
                if tag.name == "h2":
                    doc.add_heading(text, level=1)
                elif tag.name == "h3":
                    doc.add_heading(text, level=2)
                elif tag.name == "h4":
                    doc.add_heading(text, level=3)
                elif tag.name == "li":
                    doc.add_paragraph(f"• {text}")
                else:
                    doc.add_paragraph(text)

            doc.save(filePath)
            self._handle_cumulative_log_message(f"Export DOCX (basic): saved to {filePath}")
            QMessageBox.information(self, "Export Successful", f"Report saved to:\n{filePath}")
        except Exception as e:
            self._handle_cumulative_log_message(f"Basic DOCX export failed: {e}")
            QMessageBox.critical(self, "Export Error", f"Could not save DOCX report:\n{e}")

    @pyqtSlot()
    def _save_log_to_file(self):
        log_content = self.log_monitor_browser.toPlainText()
        if not log_content: QMessageBox.information(self, "Empty Log", "There is no log content to save."); return
        default_filename = f"AnalysisLog_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}.txt"
        filePath, _ = QFileDialog.getSaveFileName(self, "Save Log File", default_filename, "Log Files (*.log *.txt)")
        if filePath:
            try:
                with open(filePath, 'w', encoding='utf-8') as f:
                    f.write(log_content)
                QMessageBox.information(self, "Log Saved", f"Log successfully saved to:\n{filePath}")
            except Exception as e:
                QMessageBox.critical(self, "Save Log Error", f"Could not save log file:\n{e}")

    def _generate_rich_html_for_ai_object(self, ai_step_package: dict, base_title: str) -> str:
        """
        Builds a comprehensive HTML block for displaying AI-generated section details,
        including prompts, intermediate drafts, review feedback, and fallback information.
        It expects ai_step_package to be the direct output from an AI service function,
        which typically has a "data" key containing the payload.
        """
        is_dark = self.current_theme_name == "dark"

        # Define some basic inline styles that complement COMMON_MONITOR_CSS
        # These are derived from your theme logic in apply_theme for consistency
        fg_text_color = "#dcdcdc" if is_dark else "#1e1e1e"
        detail_item_style = "margin-bottom: 8px; padding: 10px; border-radius: 4px; line-height: 1.5;"  # Added line-height
        prompt_style = f"background-color: {'#2d333b' if is_dark else '#f6f8fa'}; border: 1px solid {'#444c56' if is_dark else '#d0d7de'}; color: {fg_text_color}; {detail_item_style} white-space: pre-wrap; word-wrap: break-word;"
        feedback_style = f"background-color: {'#3a3f2d' if is_dark else '#f0f3e4'}; border: 1px solid {'#5a5e36' if is_dark else '#e0e5c1'}; {detail_item_style}"
        fallback_style = f"background-color: {'#4d4028' if is_dark else '#fff9e6'}; border: 1px solid {'#8c6d00' if is_dark else '#ffe58f'}; color: {'#ffd700' if is_dark else '#8c6d00'}; {detail_item_style}"
        final_output_style = f"border-left: 4px solid {'#528bff' if is_dark else '#1E88E5'}; padding-left: 15px; margin-top: 10px;"

        html_parts = [f"<h3>{html.escape(base_title)}</h3>"]

        # ai_step_package IS the direct output from AI service functions like generate_ai_abstract or generate_section_with_ai_review
        # These functions return a dict like: {"type": "html_section", "data": { PAYLOAD }, "description": "..."}
        # The actual content (prompt, response, etc.) is in the "data" sub-dictionary.

        data_payload = ai_step_package.get("data", {}) if isinstance(ai_step_package, dict) else {}

        if not isinstance(data_payload, dict):  # If data_payload itself isn't a dict, something is wrong
            html_parts.append(
                f"<p class='placeholder-text'>AI output data for '{html.escape(base_title)}' is not in the expected dictionary format. Displaying raw package:</p>")
            try:
                html_parts.append(
                    f"<pre style='{prompt_style}'>{html.escape(json.dumps(ai_step_package, indent=2, default=str))}</pre>")
            except:
                html_parts.append(f"<pre style='{prompt_style}'>{html.escape(str(ai_step_package))}</pre>")
            return "".join(html_parts)

        # --- Display Fallback Information ---
        fallback_info = data_payload.get("fallback_info")
        if fallback_info:
            html_parts.append(
                f"<div style='{fallback_style}'><strong>Fallback Note:</strong> {html.escape(fallback_info)}</div>")
            original_fallback_input = data_payload.get("original_fallback_input_summary")
            if original_fallback_input:
                html_parts.append(
                    f"<details><summary style='font-weight:bold; margin-bottom:5px; cursor:pointer;'>View Input Summary Used for Fallback Draft</summary>"
                    f"<pre style='{prompt_style}'>{html.escape(original_fallback_input)}</pre></details>"
                )

        # --- Initial/R0 Prompt ---
        prompt_r0 = data_payload.get("initial_draft_prompt_sent") or data_payload.get("prompt_sent")
        if prompt_r0 and prompt_r0 != "N/A (used pre-existing/fallback text for R0)" and prompt_r0 != "N/A (fallback to initial summary for R0)" and prompt_r0 != "Prompt config error for period narrative.":
            html_parts.append(
                f"<details><summary style='font-weight:bold; margin-bottom:5px; cursor:pointer;'>View Initial Draft (R0) Prompt</summary>"
                f"<pre style='{prompt_style}'>{html.escape(prompt_r0)}</pre></details>"
            )

        # --- Initial/R0 AI Draft (show if R1/R2 also ran, implies it's not the final output) ---
        initial_draft_html = data_payload.get("initial_draft_html_response")
        if initial_draft_html and (data_payload.get("review_r1_data_used") or data_payload.get("prompt_sent_r2")):
            html_parts.append(
                f"<details><summary style='font-weight:bold; margin-bottom:5px; cursor:pointer;'>View Initial AI Draft (R0)</summary>"
                f"<div style='{detail_item_style} border-left: 3px solid #777;'>{initial_draft_html}</div></details>"
            )

        # --- Review Round 1 (R1) Data ---
        r1_data = data_payload.get("review_r1_data_used")
        if isinstance(r1_data, dict):
            r1_html_parts = [
                f"<details open><summary style='font-weight:bold; margin-bottom:5px; cursor:pointer;'>Review Round 1 Feedback</summary><div style='{feedback_style}'>"]
            has_r1_content = False
            if isinstance(r1_data.get("suggested_revisions"), list) and r1_data["suggested_revisions"]:
                has_r1_content = True
                r1_html_parts.append("<strong>Suggested Revisions:</strong><ul>")
                for rev in r1_data["suggested_revisions"]: r1_html_parts.append(f"<li>{html.escape(rev)}</li>")
                r1_html_parts.append("</ul>")
            if isinstance(r1_data.get("clarifying_questions"), list) and r1_data["clarifying_questions"]:
                has_r1_content = True
                r1_html_parts.append("<strong>Clarifying Questions:</strong><ul>")
                for q in r1_data["clarifying_questions"]: r1_html_parts.append(f"<li>{html.escape(q)}</li>")
                r1_html_parts.append("</ul>")
            search_kws = r1_data.get("search_keywords_phrases") or r1_data.get("improved_keywords_phrases")
            if isinstance(search_kws, list) and search_kws:
                has_r1_content = True
                r1_html_parts.append("<strong>Suggested Search Keywords/Phrases (for Notes):</strong><ul>")
                for kw in search_kws: r1_html_parts.append(f"<li>{html.escape(kw)}</li>")
                r1_html_parts.append("</ul>")
            if not has_r1_content: r1_html_parts.append("<p><i>No specific feedback items found in R1 data.</i></p>")
            r1_html_parts.append("</div></details>")
            html_parts.extend(r1_html_parts)

        # --- Extracted Notes Count for R2 ---
        notes_count = data_payload.get("extracted_notes_count_for_r2")
        if notes_count is not None:
            html_parts.append(f"<p><i>Number of extracted notes provided for Round 2 revision: {notes_count}</i></p>")

        # --- R2 Prompt (if R2 ran) ---
        prompt_r2 = data_payload.get("prompt_sent_r2")
        if prompt_r2:
            html_parts.append(
                f"<details><summary style='font-weight:bold; margin-bottom:5px; cursor:pointer;'>View Revision (R2) Prompt</summary>"
                f"<pre style='{prompt_style}'>{html.escape(prompt_r2)}</pre></details>"
            )

        # --- Final AI Output HTML (R2 draft, or R0 if R1/R2 were skipped) ---
        html_parts.append(f"<h4 style='margin-top:15px;'>Final Generated Section Content:</h4>")
        response_html_content = (
                data_payload.get("response_html") or
                data_payload.get("initial_draft_html_response") or
                data_payload.get("initial_draft_text")
        )
        if isinstance(response_html_content, str) and response_html_content.strip():
            html_parts.append(f"<div style='{detail_item_style}'>{response_html_content}</div>")
        else:
            html_parts.append(
                "<p style='color:red;'><i>Final AI response content not available.</i></p>")
        # --- Raw AI Response Object (from the last AI call in the sequence) ---
        raw_ai_obj_for_display = data_payload.get("ai_response_raw_r2") or \
                                 data_payload.get("ai_response_full_dict") or \
                                 data_payload.get("ai_response_object") or \
                                 data_payload.get("ai_response_raw")

        if raw_ai_obj_for_display:
            try:
                raw_display_str = json.dumps(raw_ai_obj_for_display, indent=2, default=str)
            except Exception:
                raw_display_str = str(raw_ai_obj_for_display)
            html_parts.append(
                f"<details><summary style='font-weight:bold; margin-bottom:5px; cursor:pointer;'>View Raw AI Response Object (Last Stage)</summary>"
                f"<pre style='{prompt_style}'>{html.escape(raw_display_str)}</pre></details>"
            )

        return "".join(html_parts)

    @pyqtSlot()
    def _on_monitor_prev_page(self):
        if self.current_output_view_index > 0:
            self.current_output_view_index -= 1
            self._render_current_output_view()
        self._update_monitor_nav_buttons()

    @pyqtSlot()
    def _on_monitor_next_page(self):
        if self.current_output_view_index < len(self.current_output_views) - 1:
            self.current_output_view_index += 1
            self._render_current_output_view()
        self._update_monitor_nav_buttons()

    def _update_monitor_nav_buttons(self):
        num_views = len(self.current_output_views)
        show_nav = num_views > 1
        self.monitor_prev_page_btn.setVisible(show_nav)
        self.monitor_next_page_btn.setVisible(show_nav)
        if not show_nav:
            self.monitor_prev_page_btn.setEnabled(False)
            self.monitor_next_page_btn.setEnabled(False)
            return
        self.monitor_prev_page_btn.setEnabled(self.current_output_view_index > 0)
        self.monitor_next_page_btn.setEnabled(self.current_output_view_index < num_views - 1)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    logging.basicConfig(stream=sys.stdout, level=logging.DEBUG,
                        format='%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d %(funcName)s] - %(message)s')
    try:
        from dotenv import load_dotenv

        dotenv_path = Path(__file__).resolve().parent.parent.parent / '.env'
        if dotenv_path.exists():
            load_dotenv(dotenv_path=dotenv_path); logging.info(f"Attempted to load .env from {dotenv_path}")
        else:
            logging.warning(f".env file not found at {dotenv_path}.")
    except ImportError:
        logging.warning("python-dotenv library not found.")
    widget = BibliographicAnalysisWidget();
    widget.setWindowTitle("Enhanced Bibliographic Analysis Tool");
    widget.setGeometry(50, 50, 1450, 980);
    widget.show()
    sys.exit(app.exec())